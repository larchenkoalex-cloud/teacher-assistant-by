from pathlib import Path
from typing import List

try:
    import docx  # type: ignore
except Exception:
    docx = None

try:
    from PyPDF2 import PdfReader  # type: ignore
except Exception:
    PdfReader = None


def extract_topics_from_docx(path: Path, max_topics: int = 20) -> List[str]:
    """Пытается вытащить темы из DOCX по заголовкам и коротким абзацам."""

    if docx is None:
        return []

    topics: List[str] = []
    doc = docx.Document(path)

    def _collect_from_paragraphs(paragraphs) -> None:
        for para in paragraphs:
            style_name = getattr(para.style, "name", "")
            text = para.text.strip()
            if not text:
                continue
            if style_name and "Heading" in style_name:
                topics.append(text)
            elif len(text) < 80 and text.endswith(":"):
                topics.append(text.rstrip(":"))

    # 1) Обычные абзацы
    _collect_from_paragraphs(doc.paragraphs)

    # 2) Текст внутри таблиц (часто КТП оформлены таблицами)
    for table in getattr(doc, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                _collect_from_paragraphs(cell.paragraphs)

    if not topics:
        # Fallback: короткие абзацы (в том числе из таблиц)
        def _collect_short(paragraphs) -> None:
            for para in paragraphs:
                t = para.text.strip()
                if t and 2 < len(t) < 60:
                    topics.append(t)

        _collect_short(doc.paragraphs)
        for table in getattr(doc, "tables", []):
            for row in table.rows:
                for cell in row.cells:
                    _collect_short(cell.paragraphs)
    # уникальные, максимум max_topics
    seen = set()
    result: List[str] = []
    for t in topics:
        if t not in seen:
            seen.add(t)
            result.append(t)
        if len(result) >= max_topics:
            break
    return result


def extract_topics_from_pdf(path: Path, max_topics: int = 20) -> List[str]:
    """Пытается вытащить темы из PDF по первым строкам страниц."""

    if PdfReader is None:
        return []

    reader = PdfReader(str(path))
    topics: List[str] = []
    for page in reader.pages[: max_topics]:
        text = page.extract_text() or ""
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        if lines:
            topics.append(lines[0])
    # уникальные
    seen = set()
    result: List[str] = []
    for t in topics:
        if t not in seen:
            seen.add(t)
            result.append(t)
    return result


def extract_topics_from_pdf_advanced(path: Path, max_topics: int = 50, column: int | None = None, include_text: bool = False) -> List[str]:
    """Advanced PDF topic extractor.

    Попытки (в этом порядке):
    - `pdfplumber` и извлечение таблиц (если есть). Если передан `column` — берем указанную колонку (1-based).
    - Fallback на `PyPDF2` (extract_text) — извлечение по строкам и простые эвристики.
    - OCR fallback (pdf2image + pytesseract) — только если предыдущие шаги не дали результатов.

    Возвращает список уникальных коротких строк (тем), в порядке появления.
    """
    topics: List[str] = []
    try:
        import re
        try:
            import pdfplumber
        except Exception:
            pdfplumber = None
        def _norm(s: str) -> str:
            return s.strip().replace('\xa0', ' ')

        # проверка, выглядит ли строка как тема
        def _looks_like_topic(s: str) -> bool:
            s = _norm(s)
            if not s:
                return False
            # убрать лидирующую нумерацию/маркировку
            s = re.sub(r'^\s*(?:\(?[\dА-Яа-яA-Za-z]+\)?[\.|\)]\s*|[IVXLC]+[\.|\)]\s*)', '', s)
            # явно мусор
            lower = s.lower()
            if 'http' in lower or '@' in lower or 'www.' in lower or '\\' in s:
                return False
            bad_tokens = ['содержание', 'страниц', 'страница', 'примечание', 'приложение', 'isbn', 'учебник', 'адрес', 'телефон']
            for bt in bad_tokens:
                if bt in lower:
                    return False
            # слишком много специальных символов
            non_letters = len([c for c in s if not c.isalnum() and not c.isspace()])
            if non_letters / max(1, len(s)) > 0.2:
                return False
            digits = len([c for c in s if c.isdigit()])
            if digits / max(1, len(s)) > 0.35:
                return False
            # слова и длина
            words = re.findall(r"[\w\-']+", s)
            if len(words) < 2:
                return False
            if not re.search(r'[А-Яа-яA-Za-z]', s):
                return False
            if len(s) < 6 or len(s) > 180:
                return False
            return True

        def _clean_topic(s: str) -> str:
            s = _norm(s)
            s = re.sub(r'^\s*(?:\d+\.|\(?[a-zA-Zа-яА-Я0-9]+\)\.|\(?[a-zA-Zа-яА-Я0-9]+\)\s)\s*', '', s)
            s = re.sub(r'^[\-\u2022\*\)\(]+', '', s)
            s = s.strip(' .:-–—')
            return s

        def _pick_best_cell_text(raw: str) -> str:
            """Разбить сырую строку ячейки на фрагменты и выбрать лучший, чистый фрагмент.

            Убирает хвостовые числа (артикулы/колонки), сокращает шум и возвращает
            первый подходящий фрагмент или пустую строку.
            """
            if not raw:
                return ""
            # заранее очистим специфические разделители
            parts = re.split(r'[\n\r\|;/—–]+', raw)
            candidates: list[tuple[int, str]] = []
            for p in parts:
                p_clean = _clean_topic(p)
                # убрать окончание вида ' ... 1' или ' ... 33' и лишние единичные цифры
                p_clean = re.sub(r'\s+\d{1,3}$', '', p_clean).strip()
                # иногда соседняя колонка прилеплена через несколько пробелов
                subparts = [sp.strip() for sp in re.split(r'\s{2,}', p_clean) if sp.strip()]
                if subparts:
                    for sp in subparts:
                        sp2 = sp
                        sp2 = re.sub(r'^[\-\u2022\*]+', '', sp2).strip()
                        if _looks_like_topic(sp2):
                            score = len(re.findall(r'[А-Яа-яA-Za-z]', sp2))
                            candidates.append((score, sp2))
                else:
                    if _looks_like_topic(p_clean):
                        score = len(re.findall(r'[А-Яа-яA-Za-z]', p_clean))
                        candidates.append((score, p_clean))
            if candidates:
                # возвращаем кандидат с наибольшим количеством букв (предпочитаем 'человеческий' текст)
                candidates.sort(key=lambda x: (-x[0], len(x[1])))
                return candidates[0][1]
            # если ничего не подошло — попробуем агрессивнее: уберём цифры и проверим
            fallback = re.sub(r'[\d]+', '', _clean_topic(raw)).strip()
            fallback = re.sub(r'\s{2,}', ' ', fallback)
            if _looks_like_topic(fallback):
                return fallback
            return ""

        # 1) pdfplumber + таблицы
        table_detected = False
        if pdfplumber is not None:
            try:
                with pdfplumber.open(str(path)) as pdf:
                    found_tables = False
                    # собираем все таблицы по страницам
                    all_tables = []
                    for pg in pdf.pages:
                        try:
                            tables = pg.extract_tables()
                        except Exception:
                            tables = None
                        if tables:
                            found_tables = True
                            for table in tables:
                                all_tables.append(table)
                    # если есть таблицы — обрабатываем их и пытаемся определить подходящую колонку
                    if found_tables and all_tables:
                        table_detected = True
                        # если задана колонка — используем её для всех таблиц, если возможна
                        if column and column > 0:
                            for table in all_tables:
                                for row in table:
                                    if not row:
                                        continue
                                    if 1 <= column <= len(row):
                                        cell = row[column - 1]
                                        if cell:
                                            s = _pick_best_cell_text(cell)
                                            if s:
                                                topics.append(s)
                        else:
                            # автоопределение колонки: для каждой таблицы считаем score per column
                            for table in all_tables:
                                # нормируем ширину строк
                                max_cols = max((len(r) for r in table if r), default=0)
                                if max_cols == 0:
                                    continue
                                col_scores = [0.0] * max_cols
                                col_counts = [0] * max_cols
                                for row in table:
                                    if not row:
                                        continue
                                    for ci in range(max_cols):
                                        cell = row[ci] if ci < len(row) else None
                                        if cell:
                                            s = _clean_topic(cell)
                                            col_counts[ci] += 1
                                            if _looks_like_topic(s):
                                                col_scores[ci] += 1
                                # compute ratio
                                col_ratios = [ (col_scores[i] / col_counts[i]) if col_counts[i] > 0 else 0.0 for i in range(max_cols) ]
                                # choose best column if it has reasonable ratio
                                best_idx = max(range(max_cols), key=lambda i: col_ratios[i])
                                if col_counts[best_idx] > 0 and col_ratios[best_idx] >= 0.45:
                                    # take topics from this column
                                    for row in table:
                                        if not row:
                                            continue
                                        if best_idx < len(row):
                                            s = _pick_best_cell_text(row[best_idx])
                                            if s:
                                                topics.append(s)
                                else:
                                    # если ни одна колонка не выделяется — соберём все подходящие ячейки из таблицы
                                    for row in table:
                                        if not row:
                                            continue
                                        for cell in row:
                                            if cell:
                                                s = _pick_best_cell_text(cell)
                                                if s:
                                                    topics.append(s)
                    # используем текст страницы только если явно разрешено или таблицы не найдены
                    for pg in pdf.pages:
                        try:
                            text = pg.extract_text() or ""
                            if include_text and text:
                                for ln in (text or "").splitlines():
                                    if ln and len(ln.strip()) < 200:
                                        s = _clean_topic(ln)
                                        if _looks_like_topic(s):
                                            topics.append(s)
                        except Exception:
                            pass
            except Exception:
                pass

        # 2) fallback PyPDF2 (используем если разрешено include_text или таблиц не было найдено)
        if (not topics and not table_detected) or include_text:
            try:
                if PdfReader is not None:
                    reader = PdfReader(str(path))
                    for page in reader.pages[: max_topics]:
                        try:
                            text = page.extract_text() or ""
                        except Exception:
                            text = ""
                        if text:
                            for ln in (text or "").splitlines():
                                if ln:
                                    s = _clean_topic(ln)
                                    if _looks_like_topic(s):
                                        topics.append(s)
            except Exception:
                pass

        # 3) OCR fallback — только при полном отсутствии результатов
        if not topics:
            try:
                from pdf2image import convert_from_path
                import pytesseract
                images = convert_from_path(str(path), first_page=1, last_page=min(5, max_topics))
                for img in images:
                    try:
                        text = pytesseract.image_to_string(img, lang='rus+eng')
                        for ln in text.splitlines():
                            if ln:
                                _collect_candidate(ln)
                    except Exception:
                        continue
            except Exception:
                pass

    except Exception:
        return []

    # dedupe preserving order
    seen = set()
    out: List[str] = []
    for t in topics:
        if t not in seen:
            seen.add(t)
            out.append(t)
        if len(out) >= max_topics:
            break
    return out


def extract_topics(path: Path, max_topics: int = 20) -> List[str]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_topics_from_docx(path, max_topics=max_topics)
    if suffix == ".pdf":
        try:
            return extract_topics_from_pdf_advanced(path, max_topics=max_topics)
        except Exception:
            return extract_topics_from_pdf(path, max_topics=max_topics)
    return []


def extract_full_text(path: Path) -> List[str]:
    """Возвращает список строк полного текста документа (docx / pdf).

    Удобно для отображения и ручной пометки строк как тем.
    """
    suffix = path.suffix.lower()
    lines: List[str] = []
    if suffix == ".docx":
        if docx is None:
            return []
        doc = docx.Document(path)

        # Сначала абзацы в порядке документа
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                lines.append(t)

        # Затем содержимое таблиц (если есть)
        for table in getattr(doc, "tables", []):
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        t = para.text.strip()
                        if t:
                            lines.append(t)

    elif suffix == ".pdf":
        if PdfReader is None:
            return []
        reader = PdfReader(str(path))
        for page in reader.pages:
            text = page.extract_text() or ""
            for ln in text.splitlines():
                t = ln.strip()
                if t:
                    lines.append(t)

    return lines
