import os
import re
import html
import time
import uuid
import logging
from io import BytesIO
from datetime import datetime, timezone
from pathlib import Path

import requests
import streamlit as st
import streamlit.components.v1 as components
from bs4 import BeautifulSoup, NavigableString

try:
    from streamlit_quill import st_quill
except Exception:  # pragma: no cover
    st_quill = None

from teacher_assistant.services import (
    LessonPlan,
    create_lesson_plan,
    init_db,
    list_lesson_plans,
    create_user,
    get_user_by_username,
    list_materials,
    delete_material,
    create_material,
)
from teacher_assistant.services.calendar import get_calendar_topics_for_subject
from passlib.hash import bcrypt

from parsers import extract_topics, extract_full_text, extract_topics_from_pdf_advanced
import json
from text_normalizer import normalize_ai_markdown
from markdown_utils import markdown_to_html

# Переезд логики в пакет `teacher_assistant/` (для масштабирования проекта)
from teacher_assistant.ai import openrouter as openrouter_client
from teacher_assistant.app.streamlit_helpers import safe_rerun
from teacher_assistant.utils import quill_html as quill_html_utils

try:
    from teacher_assistant.components.quill_editor import quill_editor
except Exception:  # pragma: no cover
    quill_editor = None


def generate_with_deepseek(api_key: str, prompt: str, model: str = "deepseek/deepseek-chat") -> dict:
    """Универсальная функция для запросов к DeepSeek через OpenRouter.

    Добавлены простые повторы и увеличенный таймаут, чтобы снизить вероятность
    ошибки Read timed out при временных неполадках сети или у сервиса.
    """
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://teacher-assistant.streamlit.app",
        "X-Title": "Teacher Assistant",
    }

    data = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.7,
        "max_tokens": 2000,
    }

    timeouts = [60, 90, 120]
    last_exc = None
    # unique visitor recording removed
    for to in timeouts:
        try:
            resp = requests.post(
                "https://openrouter.ai/api/v1/chat/completions",
                headers=headers,
                json=data,
                timeout=to,
            )
            resp.raise_for_status()
            return resp.json()
        except Exception as e:
            last_exc = e
            # лёгкая пауза перед следующей попыткой
            try:
                time.sleep(1 + (to // 30))
            except Exception:
                pass

    # если не удалось получить ответ
    return {"error": str(last_exc), "choices": [{"message": {"content": f"Ошибка: {last_exc}"}}]}


IS_DEV_ADMIN = False  # выключаем режим разработки по умолчанию

SUBJECTS = [
    "Математика",
    "Русский язык",
    "Литература",
    "Музыка",
    "Белорусский язык",
    "Английский язык",
    "Информатика",
    "Физика",
    "Химия",
    "Биология",
    "История",
    "Обществоведение",
    "География",
]

GRADES = [f"{i} класс" for i in range(1, 12)]

st.set_page_config(
    page_title="Teacher Assistant",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# Streamlit может запоминать состояние сайдбара в браузере и игнорировать
# initial_sidebar_state на следующих запусках. Этот хак пытается свернуть
# сайдбар на старте, проверяя фактическую видимость панели.
components.html(
        """
        <script>
        (function () {
            const collapseButtonSelectors = [
                'button[data-testid="stSidebarCollapseButton"]',
                'button[aria-label="Close sidebar"]',
                'button[title="Close sidebar"]',
                'button[aria-label="Закрыть боковую панель"]',
                'button[title="Закрыть боковую панель"]'
            ];

            function getDoc() {
                try {
                    return (window.parent && window.parent.document) ? window.parent.document : document;
                } catch (e) {
                    return document;
                }
            }

            function isSidebarVisible(doc) {
                try {
                    const sidebar = doc.querySelector('[data-testid="stSidebar"]');
                    if (!sidebar) return false;
                    const rect = sidebar.getBoundingClientRect();
                    return rect && rect.width > 40;
                } catch (e) {
                    return false;
                }
            }

            function findCollapseButton(doc) {
                for (const sel of collapseButtonSelectors) {
                    const btn = doc.querySelector(sel);
                    if (btn) return btn;
                }
                // fallback: ищем любую кнопку, похожую по aria-label/title
                try {
                    const buttons = Array.from(doc.querySelectorAll('button'));
                    const candidates = buttons.filter(b => {
                        const t = (b.getAttribute('aria-label') || b.getAttribute('title') || '').toLowerCase();
                        return t.includes('close sidebar') || t.includes('закрыть') || t.includes('боковую панель');
                    });
                    return candidates[0] || null;
                } catch (e) {
                    return null;
                }
            }

            let attempts = 0;
            const maxAttempts = 80; // ~12 секунд

            function tick() {
                attempts += 1;
                const doc = getDoc();
                if (isSidebarVisible(doc)) {
                    const btn = findCollapseButton(doc);
                    if (btn) {
                        btn.click();
                        return;
                    }
                } else {
                    // Уже свернут — ничего не делаем.
                    return;
                }

                if (attempts < maxAttempts) {
                    setTimeout(tick, 150);
                }
            }

            setTimeout(tick, 150);
            })();
            </script>
            """,
            height=0,
            width=0,
    )

# Убрана логика уникальных посетителей — оставляем только счётчик открытий страниц.

# Уменьшаем левый отступ основного контейнера (чтобы сократить расстояние до сайдбара/колонки регистрации)
# Значение уменьшено примерно вдвое относительно дефолтного.
# Добавляем стиль для узкого окна потоковой генерации, чтобы оно было фиксировано по высоте
# и не сдвигало остальные элементы страницы при поступлении новых строк.
st.markdown(
    """
    <style>
    .block-container {
        padding-left: 1rem !important;
        padding-top: 2rem !important;
        overflow: visible !important;
        position: relative !important;
        z-index: 1000 !important;
    }
    /* Предотвращаем срезание заголовка: убираем верхний margin у заголовков и поднимаем их над фоном */
    .block-container h1,
    .block-container h2,
    .block-container h3 {
        margin-top: 0 !important;
        padding-top: 0 !important;
        position: relative !important;
        z-index: 1100 !important;
    }
    /* Если какие-то фиксированные панели перекрывают контент, гарантируем видимость заголовка */
    header[data-testid="stHeader"] {
        z-index: 1200 !important;
    }
    .gen-stream-box {
        width: 100% !important;
        max-width: none !important;
        height: 220px !important;
        box-sizing: border-box !important;
        overflow: auto;
        border-radius: 6px;
        padding: 0.6rem;
        background: #fbfbfb;
        border: 1px solid #e6e6e6;
        box-shadow: none;
        font-size: 0.95rem;
        line-height: 1.4;
        white-space: pre-wrap;
        word-break: break-word;
        margin: 0;
        text-align: left;
    }
    .gen-stream-box pre {
        margin: 0;
        white-space: pre-wrap;
    }
    /* Сделать все кнопки светло-салатовыми */
    .stButton>button, div.stButton>button {
        background-color: #dff6d8 !important;
        color: #063806 !important;
        border: 1px solid #9fd48a !important;
        box-shadow: none !important;
        padding: 0.45rem 0.7rem !important;
    }
    .stButton>button:hover, div.stButton>button:hover {
        background-color: #c8f0b8 !important;
    }
    .stButton>button:active, div.stButton>button:active {
        background-color: #b0e998 !important;
    }
    .stButton>button:focus, div.stButton>button:focus {
        outline: 2px solid #8fd273 !important;
        outline-offset: 1px !important;
    }
    /* Сохранить стиль дизейбледа по умолчанию (тёмно-серый фон) */
    .stButton>button:disabled, div.stButton>button:disabled {
        opacity: 0.55 !important;
        background-color: unset !important;
        color: unset !important;
    }

    /* Увеличиваем шрифт в заголовках вкладок (st.tabs):
       это именно подписи "Раздаточный материал" / "Конспект" / "Викторина" сверху. */
    .stTabs button[role="tab"] {
        font-size: 1.2rem !important;
    }
    .stTabs button[role="tab"] p,
    .stTabs button[role="tab"] span,
    .stTabs button[role="tab"] div {
        font-size: 1.2rem !important;
        font-weight: 700 !important;
        line-height: 1.1 !important;
    }
    /* Классические стрелки у числовых полей (спиннеров) */
    input[type="number"] {
        appearance: auto !important;
        -moz-appearance: number-input !important;
    }
    input[type="number"]::-webkit-outer-spin-button,
    input[type="number"]::-webkit-inner-spin-button {
        -webkit-appearance: auto !important;
        opacity: 1 !important;
        margin: 0 !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# Инициализация БД (SQLite по умолчанию, можно заменить на Postgres через DATABASE_URL)
init_db()

import sqlite3 as _sqlite3

def _ensure_visits_table(db_path: str = "teacher_assistant_visits.db"):
    con = _sqlite3.connect(db_path, timeout=5)
    cur = con.cursor()
    cur.execute(
        """
    CREATE TABLE IF NOT EXISTS meta_visits (
        k TEXT PRIMARY KEY,
        v INTEGER,
        last_seen TEXT
    )
    """
    )
    con.commit()
    con.close()

def _ensure_clients_table(db_path: str = "teacher_assistant_visits.db"):
    # unique visitors tracking removed; keep function for compatibility
    return


def _record_client_id(client_id: str, db_path: str = "teacher_assistant_visits.db"):
    """Записать client_id напрямую в таблицу clients (без чтения из query params)."""
    if not client_id:
        return
    session_key = f"_ta_client_recorded_{client_id}"
    if st.session_state.get(session_key):
        return
    try:
        _ensure_clients_table(db_path)
        now = datetime.now(timezone.utc).isoformat()
        con = _sqlite3.connect(db_path, timeout=5)
        cur = con.cursor()
        cur.execute("SELECT visits FROM clients WHERE client_id = ?", (client_id,))
        row = cur.fetchone()
        if row:
            cur.execute(
                "UPDATE clients SET last_seen = ?, visits = visits + 1 WHERE client_id = ?",
                (now, client_id),
            )
            is_new = False
        else:
            cur.execute(
                "INSERT INTO clients (client_id, first_seen, last_seen, visits) VALUES (?, ?, ?, 1)",
                (client_id, now, now),
            )
            is_new = True
        con.commit()
        con.close()
        if is_new:
            try:
                con = _sqlite3.connect(db_path, timeout=5)
                cur = con.cursor()
                cur.execute("SELECT v FROM meta_visits WHERE k = 'unique_visitors'")
                r = cur.fetchone()
                if r:
                    cur.execute("UPDATE meta_visits SET v = v + 1, last_seen = ? WHERE k = 'unique_visitors'", (now,))
                else:
                    cur.execute("INSERT INTO meta_visits (k, v, last_seen) VALUES ('unique_visitors', 1, ?)", (now,))
                con.commit()
                con.close()
            except Exception:
                pass
        st.session_state[session_key] = True
        try:
            st.session_state["_ta_visit_unique"] = _get_unique_total(db_path)
        except Exception:
            pass
    except Exception:
        pass

def _record_client_from_query(db_path: str = "teacher_assistant_visits.db"):
    params = st.query_params
    client_id = params.get("ta_client_id", [None])[0]
    # diagnostic logging
    try:
        with open("visits_debug.log", "a", encoding="utf-8") as fh:
            fh.write(f"{datetime.now(timezone.utc).isoformat()} QUERY_PARAMS: {params}\n")
    except Exception:
        pass
    if not client_id:
        try:
            with open("visits_debug.log", "a", encoding="utf-8") as fh:
                fh.write(f"{datetime.now(timezone.utc).isoformat()} no ta_client_id in query\n")
        except Exception:
            pass
        return
    session_key = f"_ta_client_recorded_{client_id}"
    if st.session_state.get(session_key):
        return
    try:
        _ensure_clients_table(db_path)
        now = datetime.now(timezone.utc).isoformat()
        con = _sqlite3.connect(db_path, timeout=5)
        cur = con.cursor()
        cur.execute("SELECT visits FROM clients WHERE client_id = ?", (client_id,))
        row = cur.fetchone()
        if row:
            cur.execute(
                "UPDATE clients SET last_seen = ?, visits = visits + 1 WHERE client_id = ?",
                (now, client_id),
            )
            is_new = False
        else:
            cur.execute(
                "INSERT INTO clients (client_id, first_seen, last_seen, visits) VALUES (?, ?, ?, 1)",
                (client_id, now, now),
            )
            is_new = True
        con.commit()
        con.close()
        # если это новый клиент — можно опционально хранить отдельный мета-ключ
        if is_new:
            # обновим meta_visits.unique_visitors (если есть) иначе оставим подсчёт по таблице
            try:
                con = _sqlite3.connect(db_path, timeout=5)
                cur = con.cursor()
                cur.execute("SELECT v FROM meta_visits WHERE k = 'unique_visitors'")
                r = cur.fetchone()
                if r:
                    cur.execute("UPDATE meta_visits SET v = v + 1 WHERE k = 'unique_visitors'")
                else:
                    cur.execute("INSERT INTO meta_visits (k, v, last_seen) VALUES ('unique_visitors', 1, ?)", (now,))
                con.commit()
                con.close()
            except Exception:
                pass
        st.session_state[session_key] = True
        # обновим отображаемое значение уникальных посетителей
        try:
            st.session_state["_ta_visit_unique"] = _get_unique_total(db_path)
        except Exception:
            pass
    except Exception:
        pass
    # client-from-query removed — unique visitors disabled

def _get_unique_total(db_path: str = "teacher_assistant_visits.db") -> int:
    # unique visitors tracking disabled
    return 0

def _increment_page_open(db_path: str = "teacher_assistant_visits.db"):
    # Считаем только один раз за сессию рендера Streamlit
    if st.session_state.get("_ta_visit_counted"):
        return
    try:
        _ensure_visits_table(db_path)
        now = datetime.now(timezone.utc).isoformat()
        con = _sqlite3.connect(db_path, timeout=5)
        cur = con.cursor()
        cur.execute("SELECT v FROM meta_visits WHERE k = 'page_views'")
        row = cur.fetchone()
        if row:
            cur.execute(
                "UPDATE meta_visits SET v = v + 1, last_seen = ? WHERE k = 'page_views'",
                (now,),
            )
        else:
            cur.execute(
                "INSERT INTO meta_visits (k, v, last_seen) VALUES ('page_views', 1, ?)", (now,)
            )
        con.commit()
        cur.execute("SELECT v FROM meta_visits WHERE k = 'page_views'")
        total = cur.fetchone()[0]
        con.close()
        st.session_state["_ta_visit_counted"] = True
        st.session_state["_ta_visit_total"] = int(total)
    except Exception:
        # При ошибках статистики не ломаем приложение
        st.session_state["_ta_visit_total"] = st.session_state.get("_ta_visit_total", 0)

def _get_visit_total(db_path: str = "teacher_assistant_visits.db") -> int:
    try:
        _ensure_visits_table(db_path)
        con = _sqlite3.connect(db_path, timeout=5)
        cur = con.cursor()
        cur.execute("SELECT v FROM meta_visits WHERE k = 'page_views'")
        row = cur.fetchone()
        con.close()
        return int(row[0]) if row else 0
    except Exception:
        return int(st.session_state.get("_ta_visit_total", 0) or 0)

# Увеличиваем счётчик при первой загрузке пользователя в этой сессии
_increment_page_open(db_path="teacher_assistant_visits.db")

def _render_visit_counter():
    total = st.session_state.get("_ta_visit_total") or _get_visit_total()
    st.markdown(
        f"""<style>
#ta-visit-counter {{
 position: fixed;
 /* Сдвигаем влево примерно на 5 см относительно прежнего правого отступа */
 right: calc(1rem + 5cm);
 bottom: 1rem;
 background: rgba(255,255,255,0.92);
 padding: 6px 10px;
 border-radius: 8px;
 box-shadow: 0 2px 6px rgba(0,0,0,0.12);
 font-size: 0.95rem;
 z-index: 2147483647;
}}
</style>
<div id="ta-visit-counter">Открытий страницы: <strong>{total}</strong></div>""",
        unsafe_allow_html=True,
    )

# Отрисовываем счётчик (внизу справа)
_render_visit_counter()

st.title("Твой ассистет — помощник для учителя")
st.markdown("Используйте форму слева для генерации плана урока и предпросмотр справа для правок.")

# Переключение между пользовательским режимом и админ-панелью
if "is_admin_user" not in st.session_state:
    st.session_state["is_admin_user"] = False


def _has_admin_access() -> bool:
    """Лёгкая проверка админ-доступа на ранней стадии рендера.

    Нельзя вызывать is_admin() здесь, потому что она объявлена ниже по файлу.
    """

    if IS_DEV_ADMIN:
        return True
    try:
        return bool(st.session_state.get("is_admin_user"))
    except Exception:
        return False


app_mode = st.sidebar.radio("Режим", ["Пользовательский режим", "Админ-панель"], key="app_mode")


def _get_admin_password() -> str:
    """Возвращает пароль админа из secrets/env, либо дефолт.

    Важно: дефолт удобен для локальной разработки, но для продакшена
    лучше задать ADMIN_PASSWORD в st.secrets или переменных окружения.
    """

    try:
        pw = st.secrets.get("ADMIN_PASSWORD")
    except Exception:
        pw = None
    pw = pw or os.getenv("ADMIN_PASSWORD") or "adminstan"
    return str(pw)


# Защита: админ-доступ выдаётся только после ввода пароля.
# Важно: без st.stop/safe_rerun, чтобы не провоцировать ошибки фронтенда
# при резкой смене дерева элементов.
if app_mode == "Админ-панель" and not _has_admin_access():
    st.sidebar.warning("Для доступа к админ-панели нужен пароль.")
    admin_pw_input = st.sidebar.text_input("Пароль админа", type="password", key="admin_pw_input")
    if st.sidebar.button("Войти в админ-панель", key="admin_login_btn"):
        if admin_pw_input and admin_pw_input == _get_admin_password():
            st.session_state["is_admin_user"] = True
            st.sidebar.success("Админ-доступ получен")
        else:
            st.sidebar.error("Неверный пароль")

# Sidebar: показываем настройки DeepSeek только в админ-панели и только после входа
if app_mode == "Админ-панель" and _has_admin_access():
    st.sidebar.header("Настройки DeepSeek (OpenRouter)")

    # Читаем ключ OpenRouter (не DeepSeek напрямую!)
    _secrets_key = None
    try:
        _secrets_key = st.secrets.get("OPENROUTER_API_KEY")  # Изменили имя!
    except Exception:
        _secrets_key = None
    _env_key = os.getenv("OPENROUTER_API_KEY")  # Изменили имя!

    api_key_input = st.sidebar.text_input("OpenRouter API key", type="password", 
                                          help="Ключ начинается с sk-or-v1-...")
    api_key = _secrets_key or _env_key or api_key_input

    # Нормализуем: часто при копипасте добавляются пробелы/кавычки
    try:
        api_key = (api_key or "").strip().strip('"').strip("'")
    except Exception:
        pass
    if not api_key:
        api_key = None

    try:
        # Сохраняем в session_state, чтобы другие блоки (inline AI и пр.) могли читать ключ
        st.session_state["api_key"] = api_key
    except Exception:
        pass

    # Фиксированный URL для OpenRouter (не настраиваемый)
    DEEPSEEK_API_BASE = "https://openrouter.ai/api/v1/chat/completions"

    # Показываем источник ключа
    if api_key:
        if _secrets_key and api_key == _secrets_key:
            api_key_source = "st.secrets"
        elif _env_key and api_key == _env_key:
            api_key_source = "env"
        elif api_key_input and api_key == api_key_input:
            api_key_source = "sidebar"
        else:
            api_key_source = "unknown"
        st.sidebar.success(f"✅ Ключ загружен ({api_key_source})")
    else:
        st.sidebar.warning("⚠️ Ключ OpenRouter не задан")

    show_deepseek_debug = st.sidebar.checkbox("Показывать отладку API", value=False)

    if show_deepseek_debug:
        # Показываем «отпечаток» ключа (без раскрытия полного значения)
        if api_key:
            try:
                key_len = len(api_key)
                key_prefix = api_key[:10]
                key_tail = api_key[-6:]
                st.sidebar.caption(
                    f"Ключ: len={key_len}, prefix={key_prefix}…, tail=…{key_tail}, source={api_key_source}"
                )
                if not api_key.startswith("sk-or-v1-"):
                    st.sidebar.warning("Похоже, это не OpenRouter ключ. Нужен формат sk-or-v1-…")
                if api_key_source in {"st.secrets", "env"} and api_key_input:
                    st.sidebar.info(
                        "Введённый ключ в сайдбаре сейчас НЕ используется (приоритет у secrets/env). "
                        "Очистите env/secrets или задайте правильный ключ там."
                    )
            except Exception:
                pass

        try:
            with open("openrouter_debug.log", "r", encoding="utf-8") as fh:
                lines = fh.readlines()
            tail = "".join(lines[-80:])
            st.sidebar.caption("openrouter_debug.log (последние строки)")
            st.sidebar.code(tail or "(пока пусто)")
        except FileNotFoundError:
            st.sidebar.caption("openrouter_debug.log ещё не создан")
        except Exception as e:
            st.sidebar.caption(f"Не удалось прочитать openrouter_debug.log: {e}")

    # Кнопка проверки подключения
    if st.sidebar.button("Проверить подключение к API"):
        if not api_key:
            st.sidebar.error("Сначала введите API ключ")
        else:
            with st.sidebar.spinner("Проверяем..."):
                try:
                    response = requests.post(
                        DEEPSEEK_API_BASE,
                        headers={
                            "Authorization": f"Bearer {api_key}",
                            "Content-Type": "application/json",
                            "HTTP-Referer": "https://teacher-assistant.streamlit.app",
                            "X-Title": "Teacher Assistant",
                        },
                        json={
                            "model": "deepseek/deepseek-chat",
                            "messages": [{"role": "user", "content": "Ответь 'OK'"}],
                            "max_tokens": 10
                        },
                        timeout=10
                    )
                    if response.status_code == 200:
                        st.sidebar.success("✅ API работает!")
                    elif response.status_code == 401:
                        st.sidebar.error("❌ 401 Unauthorized: неверный OpenRouter API key (нужен sk-or-v1-…).")
                        try:
                            open("openrouter_debug.log", "a", encoding="utf-8").write(
                                f"[check_api] 401 unauthorized; key_source={api_key_source}\n"
                            )
                        except Exception:
                            pass
                    else:
                        st.sidebar.error(f"❌ Ошибка API: {response.status_code}")
                        if show_deepseek_debug:
                            try:
                                st.sidebar.code((response.text or "")[:2000])
                            except Exception:
                                pass
                except Exception as e:
                    st.sidebar.error(f"❌ Ошибка: {e}")


# Опция показа последних планов (по умолчанию скрыта)
show_recent_plans = st.sidebar.checkbox(
    "Показывать последние сохранённые планы",
    value=False,
    key="show_recent_plans",
)

# --- Простая аутентификация для педагогов (регистрация / вход)
if "user_id" not in st.session_state:
    st.session_state["user_id"] = None
    st.session_state["username"] = None
    st.session_state["is_admin_user"] = False

auth_mode = st.sidebar.selectbox("Аккаунт", ["Войти", "Регистрация", "Профиль"])
if auth_mode == "Регистрация":
    with st.sidebar.form("register_form"):
        reg_username = st.text_input("Логин")
        reg_email = st.text_input("Email (опционально)")
        reg_password = st.text_input("Пароль", type="password")
        reg_password2 = st.text_input("Повторите пароль", type="password")
        if st.form_submit_button("Зарегистрироваться"):
            if not reg_username or not reg_password:
                st.sidebar.warning("Укажите логин и пароль.")
            elif (reg_username or "").strip().lower() == "admin":
                st.sidebar.warning("Логин 'admin' зарезервирован для админ-панели.")
            elif reg_password != reg_password2:
                st.sidebar.warning("Пароли не совпадают.")
            elif get_user_by_username(reg_username):
                st.sidebar.warning("Пользователь с таким логином уже существует.")
            else:
                pwd_hash = bcrypt.hash(reg_password)
                user = create_user(username=reg_username, email=reg_email or None, password_hash=pwd_hash)
                st.session_state["user_id"] = user.id
                st.session_state["username"] = user.username
                st.sidebar.success("Регистрация прошла успешно. Вы вошли как {}".format(user.username))

elif auth_mode == "Войти":
    with st.sidebar.form("login_form"):
        login_username = st.text_input("Логин")
        login_password = st.text_input("Пароль", type="password")
        if st.form_submit_button("Войти"):
            # Вход в админ-панель по паролю. Пароль берётся из st.secrets.ADMIN_PASSWORD,
            # затем из env ADMIN_PASSWORD, иначе используется дефолт 'adminstan'.
            try:
                admin_pass = st.secrets.get("ADMIN_PASSWORD")
            except Exception:
                admin_pass = None
            if not admin_pass:
                admin_pass = os.getenv("ADMIN_PASSWORD") or "adminstan"

            if login_username == "admin" and login_password and login_password == admin_pass:
                st.session_state["user_id"] = 0
                st.session_state["username"] = "admin"
                st.session_state["is_admin_user"] = True
                st.sidebar.success("Вы вошли как admin (админ)")
            else:
                user = get_user_by_username(login_username)
                if not user:
                    st.sidebar.error("Пользователь не найден.")
                else:
                    try:
                        if bcrypt.verify(login_password, user.password_hash):
                            st.session_state["user_id"] = user.id
                            st.session_state["username"] = user.username
                            st.sidebar.success(f"Вы вошли как {user.username}")
                        else:
                            st.sidebar.error("Неверный пароль.")
                    except Exception:
                        st.sidebar.error("Ошибка проверки пароля.")

else:  # Профиль
    if st.session_state.get("user_id"):
        st.sidebar.write(f"Вошли как: {st.session_state.get('username')}")
        if st.sidebar.button("Выйти"):
            st.session_state["user_id"] = None
            st.session_state["username"] = None
            st.session_state["is_admin_user"] = False
            st.sidebar.success("Вы вышли.")
    else:
        st.sidebar.info("Войдите или зарегистрируйтесь, чтобы публиковать материалы.")

current_user = None
if st.session_state.get("username"):
    current_user = get_user_by_username(st.session_state["username"])

materials_dir = Path(os.getenv("MATERIALS_DIR", "materials"))
materials_dir.mkdir(exist_ok=True)


def generate_lesson_plan_locally(subject: str, grade: str, topic: str, notes: str, class_level: str = None) -> str:
    """Локальная заготовка плана урока на случай отсутствия API."""

    header = (
        f"План урока по предмету: {subject} (класс: {grade})\n"
        f"Уровень подготовки класса: {class_level or 'средний'}\n"
        f"Тема: {topic}\n\n"
    )
    body = """Цели урока:
- сформировать понимание ключевых понятий по теме;
- развивать навыки самостоятельной работы и критического мышления;
- закрепить материал через практические задания.

Ход урока:
1. Организационный момент (2–3 мин).
2. Актуализация знаний (5–7 мин).
3. Объяснение нового материала (15–20 мин).
4. Закрепление (индивидуальные и групповые задания) (10–15 мин).
5. Рефлексия и подведение итогов (5 мин).
6. Домашнее задание.

Материалы и ресурсы:
- презентация/конспект для учителя;
- раздаточные материалы для учащихся;
- дополнительные ресурсы (ссылки, видео, интерактивы).
"""
    if notes:
        body += f"\nОсобенности класса / примечания учителя:\n{notes}\n"
    return header + body


def stream_generate_chat_via_api(*, messages: list, headers: dict, placeholder, model: str = "deepseek/deepseek-chat") -> str:
    """Потоково читает ответ (SSE) и обновляет placeholder; возвращает полный текст.

    Если провайдер не поддерживает SSE, делает обычный запрос и возвращает полный ответ.
    """
    # Вынесено в teacher_assistant.ai.openrouter (для повторного использования и тестирования)
    api_key = None
    try:
        auth = (headers or {}).get("Authorization") or ""
        if auth.startswith("Bearer "):
            api_key = auth[len("Bearer ") :].strip()
    except Exception:
        api_key = None

    if not api_key:
        return ""

    def _on_update(full_text: str) -> None:
        # Рендерим поток в фиксированном окне, чтобы не "прыгала" страница.
        safe_text = html.escape(full_text or "")
        placeholder.markdown(
            f"<div class='gen-stream-box'><pre>{safe_text}</pre></div>",
            unsafe_allow_html=True,
        )

        # Подставляем примечания учителя в авто-промпт, если поле заполнено
        quiz_notes_txt = (st.session_state.get("quiz_notes") or "").strip()
        if quiz_notes_txt:
            auto_prompt += f"\nПримечания учителя: {quiz_notes_txt}\n"

    try:
        return openrouter_client.stream_chat_completions(
            api_key=api_key,
            messages=messages,
            on_update=_on_update,
            model=model,
        )
    except openrouter_client.PaymentRequiredError:
        # Пробрасываем дальше специфичную ошибку оплаты, чтобы UI мог её отобразить
        raise
    except Exception as e:
        try:
            open("openrouter_debug.log", "a", encoding="utf-8").write(
                f"[stream_generate_chat_via_api] EXCEPTION model={model!r} messages={len(messages) if messages else 0} type={type(e).__name__} msg={e!r}\n"
            )
        except Exception:
            pass
        raise


def is_admin() -> bool:
    """Проверка, является ли текущим пользователем администратором.

    В режиме разработки (IS_DEV_ADMIN=True) всегда возвращает True,
    чтобы упростить локальную работу.
    """

    if IS_DEV_ADMIN:
        return True
    # Админ определяется флагом сессии (вход по паролю ставит флаг `is_admin_user`)
    try:
        if st.session_state.get("is_admin_user"):
            return True
    except Exception:
        pass
    if current_user and getattr(current_user, "role", "user") == "admin":
        return True
    return False


def _slugify(value: str) -> str:
    """Генерация безопасного имени файла.

    Важно: поддерживаем кириллицу (Windows/macOS/Linux это обычно позволяют),
    иначе темы/предметы на русском превращаются в пустые строки.
    """

    s = (value or "").strip().lower()
    # приводим пробелы/подчёркивания к дефисам
    s = re.sub(r"[\s_]+", "-", s)
    # разрешаем латиницу, цифры, дефис/подчёркивание и кириллицу
    s = re.sub(r"[^0-9a-zа-яё\-_]+", "", s)
    # схлопываем повторяющиеся разделители
    s = re.sub(r"[-_]{2,}", "-", s)
    s = s.strip("-_")
    return (s[:80] if s else "material")



def _build_openrouter_headers(api_key: str) -> dict:
    return openrouter_client.build_headers(api_key)




def _normalize_html_for_change_detection(html_text: str) -> str:
    """Нормализует HTML для сравнения.

    Quill часто держит хвостовой `<p><br></p>` как "место для курсора".
    Его игнорируем при сравнении, чтобы не перерисовывать редактор на каждый ввод.
    """

    s = (html_text or "").strip()
    # Срезаем хвостовые пустые параграфы
    s = re.sub(
        r"(?:<p>(?:\s|&nbsp;)*<br\s*/?>(?:\s|&nbsp;)*</p>\s*)+$",
        "",
        s,
        flags=re.IGNORECASE,
    )
    s = re.sub(r"(?:<p>(?:\s|&nbsp;)*</p>\s*)+$", "", s, flags=re.IGNORECASE)
    return s.strip()

def _build_lesson_plan_messages(*, subject: str, grade: str, topic: str, lesson_type: str, class_level: str, notes: str, extra_instructions: str = "") -> list:
    # Языковые особенности — добавляем инструкции, если предмет — белорусский или английский
    language_tail = ""
    subj_lower = (subject or "").lower()
    if "белору" in subj_lower or "беларус" in subj_lower:
        language_tail = (
            "Ответ сформулируй на белорусском языке. Учитывай нормы белорусской орфографии и фонетики, "
            "включи упражнения на чтение, письмо и работу с текстом, варианты дифференцированных заданий. "
            "Адаслай на беларускай мове."
        )
    elif "иностранн" in subj_lower and "англ" in subj_lower or "английск" in subj_lower:
        language_tail = (
            "Respond in English. Focus on foreign-language teaching methods: communicative tasks, "
            "speaking/listening practice, controlled vocabulary activities, and level-appropriate (A1-A2) differentiation."
        )

    base_prompt = f"""
    СОЗДАЙ ДЕТАЛЬНЫЙ ПЛАН УРОКА ДЛЯ УЧИТЕЛЯ

    ПРЕДМЕТ: {subject}
    КЛАСС: {grade}
    УРОВЕНЬ ПОДГОТОВКИ КЛАССА: {class_level or 'средний'}
    ТИП УРОКА: {lesson_type}
    ТЕМА: {topic}
    ВРЕМЯ: 45 минут
    ДОПОЛНИТЕЛЬНО: {notes if notes else 'нет'}

    СТРУКТУРА (обязательно):
    1. Тема урока
    2. Цели урока (предметные, метапредметные, личностные)
    3. Планируемые результаты
    4. Оборудование и ресурсы
    5. Ход урока с точным таймингом:
       - Организационный момент (2-3 мин)
       - Актуализация знаний (5-7 мин)
       - Изучение нового материала (15-20 мин)
       - Закрепление (10-12 мин)
       - Рефлексия и домашнее задание (3-5 мин)
    6. Дифференцированное домашнее задание
    7. Приложения (материалы для урока)

    Формат: Markdown с заголовками ##
    Будь конкретен, практичен, учитывай ФГОС.
    ВКЛЮЧИ РАЗДЕЛЫ С ЗАДАНИЯМИ И АКТИВНОСТЯМИ:
    - Приведи 3 задания по уровням сложности (для слабого/среднего/сильного учащихся).
    - Подбери формы работы и типы упражнений, подходящие для указанного уровня подготовки класса.
    - Для смешанного класса предложи варианты дифференциации и адаптации для разноуровневых групп.
    """

    curriculum_tail = (
        f"Составь план в соответствии с учебной программой по {subject} для учреждений общего среднего образования "
        "Республики Беларусь (утверждённой Министерством образования Республики Беларусь)."
    )

    prompt = base_prompt + "\n\n" + curriculum_tail
    if language_tail:
        prompt = prompt + "\n\n" + language_tail
    if extra_instructions:
        prompt = prompt + "\n\nДополнительно: " + extra_instructions

    language_mode = None
    if "белору" in subj_lower or "беларус" in subj_lower:
        language_mode = "be"
    elif "иностранн" in subj_lower and "англ" in subj_lower or "английск" in subj_lower:
        language_mode = "en"

    system_message = None
    if language_mode == "be":
        system_message = (
            "Адказвай па-беларуску. Улічвай нормы беларускай арфаграфіі і фанетыкі; "
            "давай практычныя заданні па чытанню, пісьму і працы з тэкстам; уключай дыферэнцыраваныя варыянты."
        )
    elif language_mode == "en":
        system_message = (
            "Respond in English. Focus on foreign-language teaching methods: communicative tasks, "
            "speaking/listening practice, controlled vocabulary activities, and level-appropriate differentiation."
        )

    messages = []
    if system_message:
        messages.append({"role": "system", "content": system_message})
    messages.append({"role": "user", "content": prompt})
    return messages


def _postprocess_plan_text(raw_text: str) -> str:
    import re

    if not raw_text:
        return ""

    # Базовая нормализация переносов
    text = (
        raw_text.replace("\r\n", "\n")
        .replace("\r", "\n")
        .replace("\u200b", "")
        .replace("\xa0", "")
    ).strip("\n")

    # Разбиваем на строки
    lines = text.split("\n")

    # Убираем обёртку ```...```
    if lines and lines[0].lstrip().startswith("``"):
        for i in range(len(lines) - 1, 0, -1):
            if lines[i].strip() == "```":
                lines = lines[1:i]
                break

    # Убираем пустые пункты списков
    empty_list_item_re = re.compile(r"^\s*(([-*•])|(\d+\.))\s*$")
    lines = [ln for ln in lines if not empty_list_item_re.match(ln)]

    # Убираем общий левый отступ (если он большой)
    non_empty = [ln for ln in lines if ln.strip()]
    if non_empty:
        min_indent = min(len(ln) - len(ln.lstrip(" ")) for ln in non_empty)
        if min_indent >= 4:
            lines = [ln[min_indent:] if len(ln) >= min_indent else ln for ln in lines]

    # Функция проверки, является ли строка элементом списка
    def is_list_item(s: str) -> bool:
        s = s.strip()
        return (
            s.startswith(("-", "*", "•"))
            or re.match(r"^\d+\.", s) is not None
        )

    # 🔥 Основная очистка пустых строк
    cleaned = []
    prev_empty = False

    for i, ln in enumerate(lines):
        stripped = ln.strip()

        # Пустая строка
        if stripped == "":
            # Удаляем пустую строку между элементами списка
            if i > 0 and i < len(lines) - 1:
                if is_list_item(lines[i - 1]) and is_list_item(lines[i + 1]):
                    continue

            # Удаляем пустую строку сразу после заголовков
            if cleaned and cleaned[-1].strip().startswith(("#", "##", "###", "####")):
                continue

            # Обычная логика: не даём двум пустым подряд
            if not prev_empty:
                cleaned.append("")
            prev_empty = True
            continue

        # Непустая строка
        cleaned.append(ln)
        prev_empty = False

    # Убираем пустые строки в начале и конце
    while cleaned and cleaned[0].strip() == "":
        cleaned.pop(0)
    while cleaned and cleaned[-1].strip() == "":
        cleaned.pop()

    return "\n".join(cleaned)


def _normalize_docx_filename(title: str) -> str:
    safe = re.sub(r"[^0-9A-Za-zА-Яа-яЁё_\- ]+", "", (title or "")).strip()
    safe = safe[:80] if safe else "lesson_plan"
    return f"{_slugify(safe)}.docx"


def _normalize_material_filename(*, grade: str | None, kind: str, topic: str | None, ext: str = "docx") -> str:
    """Имя файла по шаблону: класс-(план|раздатка|викторина)-тема.

    Пример: "5-класс-план-десятичные-дроби.docx".
    """
    kind_map = {
        "план": "план",
        "конспект": "план",
        "раздатка": "раздатка",
        "раздаточный материал": "раздатка",
        "викторина": "викторина",
        "quiz": "викторина",
    }
    kind_norm = kind_map.get((kind or "").strip().lower(), (kind or "").strip().lower() or "план")

    grade_part = (grade or "").strip()
    grade_part = re.sub(r"\s+", "-", grade_part)

    topic_part = (topic or "").strip()
    topic_part = re.sub(r"\s+", "-", topic_part)

    base = "-".join([p for p in [grade_part, kind_norm, topic_part] if p])
    if not base:
        base = kind_norm or "material"

    safe = re.sub(r"[^0-9A-Za-zА-Яа-яЁё_\- ]+", "", base).strip()
    safe = safe[:120] if safe else "material"
    ext = (ext or "docx").lstrip(".")
    return f"{_slugify(safe)}.{ext}"


def _html_to_docx_bytes(html: str) -> bytes:
    from docx import Document

    def add_inline(paragraph, node, *, bold=False, italic=False):
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
            return

        if not hasattr(node, "name"):
            return

        name = (node.name or "").lower()
        if name in {"strong", "b"}:
            for child in node.children:
                add_inline(paragraph, child, bold=True or bold, italic=italic)
            return
        if name in {"em", "i"}:
            for child in node.children:
                add_inline(paragraph, child, bold=bold, italic=True or italic)
            return
        if name == "br":
            paragraph.add_run("\n")
            return

        # span/a/etc — просто рекурсивно обходим
        for child in node.children:
            add_inline(paragraph, child, bold=bold, italic=italic)

    soup = BeautifulSoup((html or ""), "html.parser")
    doc = Document()

    body = soup.body if soup.body else soup

    def add_block(tag):
        name = (tag.name or "").lower()
        if name in {"h1", "h2", "h3", "h4"}:
            level = {"h1": 1, "h2": 2, "h3": 3, "h4": 4}[name]
            text = tag.get_text(" ", strip=True)
            if text:
                doc.add_heading(text, level=level)
            return
        if name in {"p", "div"}:
            # Quill часто отдаёт <p><br></p>
            if not tag.get_text(strip=True) and not tag.find(["strong", "em", "b", "i"]):
                doc.add_paragraph("")
                return
            p = doc.add_paragraph("")
            for child in tag.children:
                add_inline(p, child)
            return
        if name in {"ul", "ol"}:
            style = "List Bullet" if name == "ul" else "List Number"
            for li in tag.find_all("li", recursive=False):
                p = doc.add_paragraph("", style=style)
                for child in li.children:
                    # поддержка вложенных списков
                    if getattr(child, "name", None) and child.name.lower() in {"ul", "ol"}:
                        add_block(child)
                    else:
                        add_inline(p, child)
            return

        if name == "table":
            # Собираем строки tr (включая thead/tbody)
            rows = []
            for tr in tag.find_all("tr", recursive=False):
                rows.append(tr)
            # Иногда таблица имеет thead/tbody
            if not rows:
                for section in tag.find_all(["thead", "tbody"], recursive=False):
                    for tr in section.find_all("tr", recursive=False):
                        rows.append(tr)

            if not rows:
                return

            # Определим число колонок по максимуму
            max_cols = 0
            table_cells = []
            for tr in rows:
                cells = tr.find_all(["th", "td"], recursive=False)
                table_cells.append(cells)
                if len(cells) > max_cols:
                    max_cols = len(cells)

            if max_cols == 0:
                return

            doc_table = doc.add_table(rows=len(rows), cols=max_cols)
            try:
                doc_table.style = "Table Grid"
            except Exception:
                pass

            for r_idx, cells in enumerate(table_cells):
                for c_idx in range(max_cols):
                    cell_tag = cells[c_idx] if c_idx < len(cells) else None
                    cell = doc_table.cell(r_idx, c_idx)
                    # Очистим параграф по умолчанию
                    cell.text = ""
                    p = cell.paragraphs[0]
                    if cell_tag is None:
                        continue
                    # Если это заголовочная ячейка <th>, делаем текст полужирным
                    is_header = (cell_tag.name or "").lower() == "th"
                    for child in cell_tag.children:
                        if getattr(child, "name", None) and child.name.lower() in {"ul", "ol"}:
                            # вложенный список — добавим как обычный текст с переносами
                            for li in child.find_all("li"):
                                add_inline(p, li)
                                p.add_run("\n")
                        else:
                            add_inline(p, child, bold=is_header)
            return

        # неизвестный блок — пытаемся обработать детей
        for child in tag.children:
            if getattr(child, "name", None):
                add_block(child)

    for child in body.children:
        if getattr(child, "name", None):
            add_block(child)

    buff = BytesIO()
    doc.save(buff)
    return buff.getvalue()




# Разделяем экран: левая колонка — форма генерации, правая — редактор результата
# Уменьшаем левый отступ/ширину левой колонки вдвое (раньше было [2,3])
col_form, col_editor = st.columns([1, 3])

with col_form:
    st.markdown("<h3 style=\"text-align:left; margin-top:0.25rem; margin-bottom:0.5rem;\">Генерация плана урока (ИИ)</h3>", unsafe_allow_html=True)
    def _gen_meta_changed():
        # При изменении верхних метаданных синхронизируем их с областью раздатки
        try:
            st.session_state["handout_subject"] = st.session_state.get("gen_subject") or SUBJECTS[0]
            st.session_state["handout_grade"] = st.session_state.get("gen_grade") or GRADES[0]
            # Тему синхронизируем только если поле пустое или совпадает со старой темой
            st.session_state["handout_topic"] = st.session_state.get("gen_topic") or ""
            # Пометим раздатку как требующую регенерации
            st.session_state["handout_generated_md"] = ""
            st.session_state["handout_prompt_autofill"] = True
        except Exception:
            pass

    grade = st.selectbox("Класс", GRADES, index=3, key="gen_grade", on_change=_gen_meta_changed)
    subject = st.selectbox("Предмет", SUBJECTS, index=0, key="gen_subject", on_change=_gen_meta_changed)
    topic = st.text_area(
        "Тема урока",
        value=st.session_state.get("gen_topic", ""),
        placeholder="Десятичные дроби",
        help="Ключевая тема занятия (можно в несколько строк)",
        key="gen_topic",
        height=50,
        on_change=_gen_meta_changed,
    )
    # Тип урока — выбирает учитель
    lesson_type = st.selectbox(
        "Тип урока",
        [
            "объяснение нового материала",
            "комбинированный",
            "закрепление",
            "повторение",
            "контрольный",
            "практическая работа / лабораторная",
        ],
        index=1,
    )
    class_level = st.selectbox(
        "Уровень подготовки класса",
        ["слабый", "средний", "сильный", "смешанный"],
        index=1,
        key="gen_class_level",
    )
    notes = st.text_area("Особенности класса / пожелания", placeholder="Особенность класса, акценты, приём мотивации, что важно подчеркнуть...")
    # (Упрощено) шаблон промпта можно править вручную внизу — кнопка автозагрузки убрана

    # Редактируемый шаблон промпта (можно подправить перед генерацией)
    if "prompt_template" not in st.session_state:
        st.session_state["prompt_template"] = ""
    # Кнопка генерации перемещена над шаблоном промпта по просьбе пользователя
    generate_clicked = st.button("Сгенерировать план урока")
    # Подтверждение очистки перед генерацией нового плана
    if "confirm_new_generation" not in st.session_state:
        st.session_state["confirm_new_generation"] = False
    if False:
        with st.expander("Шаблон запроса (можно править)"):
            btns = st.columns([1, 1])
            if btns[0].button("Загрузить запрос по умолчанию"):
                if not subject or not topic:
                    st.warning("Укажите предмет и тему, чтобы загрузить автопромпт.")
                else:
                    try:
                        msgs = _build_lesson_plan_messages(
                            subject=subject,
                            grade=grade,
                            topic=topic,
                            lesson_type=lesson_type,
                            class_level=class_level,
                            notes=notes,
                            extra_instructions="",
                        )
                        if msgs:
                            st.session_state["prompt_template"] = msgs[-1].get("content", "")
                    except Exception as e:
                        st.error(f"Не удалось сгенерировать промпт: {e}")
            if btns[1].button("Сбросить запрос"):
                st.session_state["prompt_template"] = ""
            st.text_area("Шаблон запроса", key="prompt_template", height=200, placeholder="Оставьте пустым для автоподстановки")
    # Скрываем селектбокс выбора источника генерации — фиксируем Deepseek по умолчанию
    model_choice = "Deepseek API (через ключ)"
    # Временно убираем выбор видимости из UI — все сохранения будут помечаться как private
    visibility = "private"

    # Подсказки тем из существующих материалов и планов для выбранного предмета и класса
    existing_plans = list_lesson_plans(limit=300)
    existing_materials = list_materials(limit=300)

    topic_candidates = set()
    for p in existing_plans:
        if p.subject == subject and p.grade == grade and p.topic:
            topic_candidates.add(p.topic)
    for m in existing_materials:
        if getattr(m, "subject", None) == subject and getattr(m, "grade", None) == grade and m.topics:
            for t in m.topics.split(","):
                t = t.strip()
                if t:
                    topic_candidates.add(t)

    # Подсадим темы из календарно-тематического планирования (KTP)
    try:
        ktp_topics = get_calendar_topics_for_subject(subject)
        for t in ktp_topics:
            topic_candidates.add(t)
    except Exception:
        pass

    typed = st.session_state.get("gen_topic", "") or ""
    if typed:
        # Общие совпадения из материалов и планов
        matches = [t for t in topic_candidates if typed.lower() in t.lower()]
        matches = sorted(matches)[:10]
        if matches:
            st.caption("Возможные темы из планов/материалов:")
            for i, t in enumerate(matches):
                if st.button(t, key=f"suggest_topic_{i}"):
                    st.session_state["gen_topic"] = t
                    safe_rerun()

        # Подсказки конкретно из календарно-тематического планирования (KTP)
        try:
            ktp_topics = get_calendar_topics_for_subject(subject)
        except Exception:
            ktp_topics = []

        ktp_matches = [t for t in ktp_topics if typed.lower() in t.lower()]
        ktp_matches = sorted(ktp_matches)[:10]
        if ktp_matches:
            st.caption("Подсказки из KTP:")
            for i, t in enumerate(ktp_matches):
                if st.button(f"{t} — из KTP", key=f"ktp_suggest_{i}"):
                    st.session_state["gen_topic"] = t
                    st.success("Тема взята из календарно‑тематического планирования.")
                    safe_rerun()

    # Проверяем, есть ли уже материал в предпросмотре
    existing_preview = bool(
        (st.session_state.get("generated_content") or "")
        or (st.session_state.get("preview_html") or "")
        or (st.session_state.get("stream_buffer") or "")
    )

    def _begin_generation() -> None:
        """Очищает предпросмотр и инициализирует состояние новой генерации."""
        st.session_state["preview_apply_replace"] = None
        st.session_state["preview_request_selection"] = None
        st.session_state["preview_request_uid"] = None
        st.session_state["preview_pending_action"] = None
        st.session_state["preview_selected_range"] = None
        st.session_state["preview_selected_text"] = ""
        st.session_state["preview_rewrite_range"] = None
        st.session_state["preview_rewrite_source"] = ""
        st.session_state["preview_rewrite_result"] = ""
        st.session_state["preview_src_view"] = ""
        st.session_state["preview_res_view"] = ""
        st.session_state["preview_fill_widgets"] = None
        st.session_state["preview_html"] = ""
        st.session_state["generated_content"] = ""

        api_key_local = st.session_state.get("api_key") or os.getenv("OPENROUTER_API_KEY")
        if not api_key_local:
            st.error("Укажите OpenRouter API key в админ-панели (sk-or-v1-...).")
            return

        headers_local = _build_openrouter_headers(api_key_local)
        messages_local = _build_lesson_plan_messages(
            subject=subject,
            grade=grade,
            topic=topic,
            lesson_type=lesson_type,
            class_level=class_level,
            notes=notes,
            extra_instructions=st.session_state.get("prompt_template", ""),
        )

        st.session_state["is_generating"] = True
        st.session_state["start_stream_now"] = True
        st.session_state["stream_buffer"] = ""
        st.session_state["generated_headers"] = headers_local
        st.session_state["generated_messages"] = messages_local
        st.session_state["generated_model"] = "deepseek/deepseek-chat"
        topic_title = (topic or "").splitlines()[0].strip() or (topic or "")
        st.session_state["generated_title"] = f"{subject or 'Урок'} — {topic_title}"[:200]

    # 1) Нажатие на генерацию
    if generate_clicked:
        if not subject or not topic:
            st.warning("Укажите хотя бы предмет и тему урока.")
        elif existing_preview:
            st.session_state["confirm_new_generation"] = True
        else:
            _begin_generation()
            safe_rerun()

    # 2) Экран подтверждения (виден до тех пор, пока пользователь не подтвердит/не отменит)
    if st.session_state.get("confirm_new_generation"):
        st.warning(
            "Сейчас в предпросмотре уже есть материал. Новый план очистит текущий текст. "
            "Подтвердите, чтобы продолжить."
        )
        c1, c2 = st.columns([1, 1])
        if c1.button("Подтвердить и сгенерировать", key="confirm_generate_btn"):
            st.session_state["confirm_new_generation"] = False
            if not subject or not topic:
                st.warning("Укажите хотя бы предмет и тему урока.")
            else:
                _begin_generation()
                safe_rerun()
        if c2.button("Отмена", key="cancel_generate_btn"):
            st.session_state["confirm_new_generation"] = False
            safe_rerun()

    # 3) Обычная логика генерации (для кейса без подтверждения)
    if False:
        pass

    if False:
        if not subject or not topic:
            st.warning("Укажите хотя бы предмет и тему урока.")
        elif model_choice == "Deepseek API (через ключ)":
            if not api_key:
                st.error("Укажите OpenRouter API key в админ-панели (sk-or-v1-...).")
            else:
                headers = _build_openrouter_headers(api_key)
                messages = _build_lesson_plan_messages(
                    subject=subject,
                    grade=grade,
                    topic=topic,
                    lesson_type=lesson_type,
                    class_level=class_level,
                    notes=notes,
                    extra_instructions=st.session_state.get("prompt_template", ""),
                )

                # Начинаем новую генерацию — сбрасываем состояние предпросмотра и буферы
                st.session_state["preview_apply_replace"] = None
                st.session_state["preview_request_selection"] = None
                st.session_state["preview_request_uid"] = None
                st.session_state["preview_pending_action"] = None
                st.session_state["preview_selected_range"] = None
                st.session_state["preview_selected_text"] = ""
                st.session_state["preview_rewrite_range"] = None
                st.session_state["preview_rewrite_source"] = ""
                st.session_state["preview_rewrite_result"] = ""
                st.session_state["preview_src_view"] = ""
                st.session_state["preview_res_view"] = ""
                st.session_state["preview_fill_widgets"] = None

                st.session_state["is_generating"] = True
                st.session_state["start_stream_now"] = True
                st.session_state["stream_buffer"] = ""
                st.session_state["generated_headers"] = headers
                st.session_state["generated_messages"] = messages
                st.session_state["generated_model"] = "deepseek/deepseek-chat"
                # Формируем читаемый заголовок из первой строки темы (без переносов)
                topic_title = (topic or "").splitlines()[0].strip() or (topic or "")
                st.session_state["generated_title"] = f"{subject or 'Урок'} — {topic_title}"[:200]
        else:
            plan_text = generate_lesson_plan_locally(subject, grade, topic, notes, class_level)
            # Локальная генерация: сохраняем в предпросмотр (без загрузки в редактор)
            topic_title = (topic or "").splitlines()[0].strip() or (topic or "")
            st.session_state["generated_title"] = f"{subject or 'Урок'} — {topic_title}"[:200]
            st.session_state["stream_buffer"] = plan_text
            st.session_state["generated_content"] = plan_text

            # Сброс состояния предпросмотра, чтобы старые фрагменты не оставались видимыми
            st.session_state["preview_apply_replace"] = None
            st.session_state["preview_request_selection"] = None
            st.session_state["preview_request_uid"] = None
            st.session_state["preview_pending_action"] = None
            st.session_state["preview_selected_range"] = None
            st.session_state["preview_selected_text"] = ""
            st.session_state["preview_rewrite_range"] = None
            st.session_state["preview_rewrite_source"] = ""
            st.session_state["preview_rewrite_result"] = ""
            st.session_state["preview_src_view"] = ""
            st.session_state["preview_res_view"] = ""
            st.session_state["preview_fill_widgets"] = None

            st.success("✅ План сгенерирован. Редактор скрыт — правьте текст через предпросмотр и AI-замену.")

with col_editor:
    # Центрируем заголовок в своей колонке
    st.markdown("<h3 style=\"text-align:left; margin-top:0.25rem; margin-bottom:0.5rem;\">Предпросмотр плана урока</h3>", unsafe_allow_html=True)

    # Оставляем только визуальный редактор (WYSIWYG), чтобы не отвлекать учителя Markdown-разметкой.

    stream_placeholder = st.empty()

    # Инициализация полей предпросмотра, чтобы блок предпросмотра можно было
    # отрендерить здесь сразу под заголовком (без зависимости от остальных блоков).
    if "preview_src_view" not in st.session_state:
        st.session_state["preview_src_view"] = ""
    if "preview_res_view" not in st.session_state:
        st.session_state["preview_res_view"] = ""
    if "preview_rewrite_range" not in st.session_state:
        st.session_state["preview_rewrite_range"] = None

    # Показываем блок предварительного просмотра перед всеми управлениями
    # По умолчанию свернут; разворачивается автоматически, если есть исходный
    # или сгенерированный результат, либо если в буфере есть данные.
    preview_has_content = bool(
        (st.session_state.get("preview_src_view") or "").strip()
        or (st.session_state.get("preview_res_view") or "").strip()
        or isinstance(st.session_state.get("preview_fill_widgets"), dict)
        or bool(st.session_state.get("preview_rewrite_result"))
    )

    # Редактор убран: работаем только с предпросмотром.

    # Если идёт генерация — показываем потоковый предпросмотр.
    # ВАЖНО: после окончания стрима сразу переключаемся на Quill в этом же прогоне,
    # без safe_rerun (иначе можно "застрять" без редактора).
    if st.session_state.get("is_generating"):
        if st.session_state.get("start_stream_now"):
            st.session_state["start_stream_now"] = False
            messages = st.session_state.pop("generated_messages", None)
            headers = st.session_state.pop("generated_headers", None)
            model_to_use = st.session_state.pop("generated_model", "deepseek/deepseek-chat")

            full_text = ""
            try:
                if messages and headers:
                    with st.spinner("🤖 Генерирую..."):
                        try:
                            full_text = stream_generate_chat_via_api(
                                messages=messages,
                                headers=headers,
                                placeholder=stream_placeholder,
                                model=model_to_use,
                            )
                        except openrouter_client.PaymentRequiredError as pay_err:
                            # Явно сообщаем пользователю о проблеме с оплатой
                            try:
                                msg = str(pay_err)
                            except Exception:
                                msg = "Проблема с оплатой на стороне OpenRouter (402)."
                            st.error(f"Проблема с оплатой OpenRouter: {msg}")
                            full_text = ""
                        except Exception as e:
                            st.error(f"Ошибка запроса к OpenRouter/DeepSeek: {e}")
                            st.caption("Подробности см. в openrouter_debug.log (в корне проекта).")
                            full_text = ""
            except Exception:
                full_text = ""

            if not full_text:
                full_text = generate_lesson_plan_locally(subject, grade, topic, notes, class_level)

            st.session_state["stream_buffer"] = full_text
            # Сохраняем сгенерированный Markdown отдельно — НЕ загружаем автоматически в редактор.
            st.session_state["generated_content"] = full_text
            # Сброс предпросмотра (пересчитается в HTML на следующем прогоне)
            st.session_state["preview_html"] = ""
            st.session_state["is_generating"] = False

        stream_buffer = st.session_state.get("stream_buffer", "")
        if stream_buffer:
            # Во время генерации показываем поток в фиксированном окне.
            # (Если SSE недоступен и текст пришёл целиком, покажем его здесь.)
            safe_text = html.escape(stream_buffer or "")
            stream_placeholder.markdown(
                f"<div class='gen-stream-box'><pre>{safe_text}</pre></div>",
                unsafe_allow_html=True,
            )
    if not st.session_state.get("is_generating"):
        # После генерации убираем потоковый Markdown, чтобы пользователь работал
        # только в Quill (иначе ПКМ будет открывать браузерное меню на обычном тексте страницы).
        stream_placeholder.empty()

        # --- Интерактивный предпросмотр с ПКМ→AI замена (без "режима редактора")
        preview_md = st.session_state.get("generated_content") or st.session_state.get("stream_buffer") or ""
        if "preview_html" not in st.session_state:
            st.session_state["preview_html"] = ""
        if "preview_apply_replace" not in st.session_state:
            st.session_state["preview_apply_replace"] = None
        if "preview_request_selection" not in st.session_state:
            st.session_state["preview_request_selection"] = None
        if "preview_request_uid" not in st.session_state:
            st.session_state["preview_request_uid"] = None
        if "preview_pending_action" not in st.session_state:
            st.session_state["preview_pending_action"] = None
        if "preview_selected_range" not in st.session_state:
            st.session_state["preview_selected_range"] = None
        if "preview_selected_text" not in st.session_state:
            st.session_state["preview_selected_text"] = ""
        if "preview_rewrite_range" not in st.session_state:
            st.session_state["preview_rewrite_range"] = None
        if "preview_rewrite_source" not in st.session_state:
            st.session_state["preview_rewrite_source"] = ""
        if "preview_rewrite_result" not in st.session_state:
            st.session_state["preview_rewrite_result"] = ""
        # Виджеты предпросмотра (важно: при наличии key значение из `value=`
        # применяется только при первом рендере, дальше живёт state виджета).
        if "preview_src_view" not in st.session_state:
            st.session_state["preview_src_view"] = ""
        if "preview_fill_widgets" not in st.session_state:
            st.session_state["preview_fill_widgets"] = None

        # Если на прошлом прогоне мы сгенерировали новый фрагмент, то применяем его
        # к ключам виджетов ДО их отрисовки.
        pending_fill = st.session_state.get("preview_fill_widgets")
        if isinstance(pending_fill, dict):
            st.session_state["preview_src_view"] = pending_fill.get("src", "") or ""
            st.session_state["preview_res_view"] = pending_fill.get("res", "") or ""
            st.session_state["preview_fill_widgets"] = None
        if "preview_res_view" not in st.session_state:
            st.session_state["preview_res_view"] = ""
        if "preview_event_log" not in st.session_state:
            st.session_state["preview_event_log"] = []
        # Диагностика предпросмотра и лог событий скрыты в UI по запросу пользователя.
        # Для отладки при необходимости можно временно вернуть expander'ы или добавить
        # опцию в боковую панель (например, st.sidebar.checkbox("Показывать диагностику")).

        if preview_md and not (st.session_state.get("preview_html") or "").strip():
            # Попробуем один раз сгенерировать HTML предпросмотра из Markdown, чтобы
            # компоненту пришёл непустой `value` и он мог эмитировать событие content.
            try:
                md2 = normalize_ai_markdown(_postprocess_plan_text(preview_md))
                st.session_state["preview_html"] = quill_html_utils.sanitize_html_for_quill(markdown_to_html(md2))
            except Exception as e:
                # Не фатально — записываем в лог для диагностики
                logs = st.session_state.get("preview_event_log", [])
                logs.insert(0, {"time": datetime.now(timezone.utc).isoformat(), "warning": f"preview_html generation failed: {e}"})
                st.session_state["preview_event_log"] = logs[:200]

        prev_instr_choice_key = "preview_ai_instr_choice"
        prev_instr_custom_key = "preview_ai_instr_custom"
        if prev_instr_choice_key not in st.session_state:
            st.session_state[prev_instr_choice_key] = "Расширить и добавить примеры"
        if prev_instr_custom_key not in st.session_state:
            st.session_state[prev_instr_custom_key] = ""

        INSTR_PRESETS = [
            "Расширить и добавить примеры",
            "Сократить и сделать яснее",
            "Упростить для учеников",
            "Сделать более официально",
            "Сделать более разговорно",
            "Исправить ошибки и улучшить стиль",
            "Переформулировать без изменения смысла",
            "Свой вариант...",
        ]

        def _current_instr() -> str:
            choice = (st.session_state.get(prev_instr_choice_key) or "").strip()
            if choice == "Свой вариант...":
                return (st.session_state.get(prev_instr_custom_key) or "").strip()
            return choice

        cols_instr = st.columns([1, 1])
        cols_instr[0].markdown("**Выделение → Переделать → (посмотреть) → Заменить**")
        cols_instr[0].selectbox("Инструкция для ИИ", INSTR_PRESETS, key=prev_instr_choice_key, label_visibility="collapsed")
        # Правый столбец используем для показа статуса/спиннера процесса переделки
        # (`cols_instr[1]` будет наполняться при запуске операции). 
        cols_instr[1].empty()
        if st.session_state.get(prev_instr_choice_key) == "Свой вариант...":
            st.text_input("Свой вариант инструкции", key=prev_instr_custom_key, placeholder="Например: сделай короче и добавь конкретику")

        # Предварительный просмотр переделанного фрагмента — теперь сразу под селектбокс инструкций
        with st.expander("Предварительный просмотр переделанного фрагмента", expanded=preview_has_content):
            if (st.session_state.get("preview_src_view") or "").strip():
                st.caption("Исходный фрагмент")
                st.text_area(
                    "Исходный фрагмент",
                    height=90,
                    disabled=True,
                    key="preview_src_view",
                    label_visibility="collapsed",
                )
            st.caption("Результат (переделанный фрагмент)")
            st.text_area(
                "Результат",
                height=140,
                key="preview_res_view",
                label_visibility="collapsed",
            )

        btn_col1, btn_col2 = st.columns([1, 1])
        if btn_col1.button("Переделать выделенный фрагмент"):
            instr = _current_instr()
            if not instr:
                st.warning("Введите инструкцию для ИИ.")
            else:
                req_uid = uuid.uuid4().hex
                st.session_state["preview_pending_action"] = "rewrite"
                st.session_state["preview_request_uid"] = req_uid
                st.session_state["preview_request_selection"] = {"_uid": req_uid}
                # Без дополнительного rerun: текущий прогон дойдёт до компонента,
                # и он получит requestSelection в этом же рендере.

        can_apply = bool(st.session_state.get("preview_rewrite_range")) and bool((st.session_state.get("preview_res_view") or "").strip())
        if btn_col2.button("Заменить выделенный фрагмент", disabled=not can_apply):
            replacement_text = (st.session_state.get("preview_res_view") or "").strip()
            st.session_state["preview_apply_replace"] = {
                "range": st.session_state.get("preview_rewrite_range"),
                "text": replacement_text,
                "_uid": time.time(),
            }
            # applyReplace будет отправлен в компонент в этом же прогоне.

        # Cloud-escape hatch: если Quill/Custom Components ломают фронтенд,
        # можно отключить компонент через Secrets/ENV: DISABLE_QUILL=1
        disable_quill_env = (os.environ.get("DISABLE_QUILL") or "").strip().lower() in ("1", "true", "yes")
        try:
            disable_quill_secret_raw = (st.secrets.get("DISABLE_QUILL") or "")
        except Exception:
            disable_quill_secret_raw = ""
        disable_quill_secret = str(disable_quill_secret_raw).strip().lower() in ("1", "true", "yes")
        disable_quill = disable_quill_env or disable_quill_secret

        evt_preview = None
        if quill_editor is not None and not disable_quill:
            try:
                evt_preview = quill_editor(
                    value=st.session_state.get("preview_html") or "",
                    height=420,
                    placeholder="Сгенерируйте план — он появится здесь...",
                    apply_replace=st.session_state.get("preview_apply_replace"),
                    request_selection=st.session_state.get("preview_request_selection"),
                    key="preview_quill",
                )
            except Exception as _err:
                st.warning("Компонент Quill вызвал ошибку — использую текстовый fallback.")
                try:
                    logs = st.session_state.get("preview_event_log", [])
                    logs.insert(0, {"time": datetime.now(timezone.utc).isoformat(), "evt": f"quill_error: {_err!r}"})
                    st.session_state["preview_event_log"] = logs[:200]
                except Exception:
                    pass
        else:
            if preview_md:
                st.info("Компонент предпросмотра (Quill) отключён/недоступен. Используется текстовый fallback.")

        if evt_preview is None:
            st.text_area(
                "preview_fallback",
                value=st.session_state.get("preview_html") or "",
                height=420,
            )

        # (экспандер перемещён выше, под селектбокс инструкций)

        if isinstance(evt_preview, dict):
            # Логируем приходящие события от компонента для диагностики
            logs = st.session_state.get("preview_event_log", [])
            logs.insert(0, {"time": datetime.now(timezone.utc).isoformat(), "evt": evt_preview})
            st.session_state["preview_event_log"] = logs[:200]
            evt_type = evt_preview.get("type")
            if evt_type == "content":
                html_value = evt_preview.get("html")
                if html_value is not None:
                    st.session_state["preview_html"] = html_value
                    # Завершаем цикл apply, если он был
                    if st.session_state.get("preview_apply_replace") is not None:
                        st.session_state["preview_apply_replace"] = None
                        # После того как компонент применил замену, сбрасываем
                        # информацию о текущем диапазоне замены, чтобы
                        # кнопка "Заменить выделенный фрагмент" стала неактивной.
                        st.session_state["preview_rewrite_range"] = None
                        # Не делаем принудительный rerun: событие content уже пришло на rerun.
            elif evt_type == "selection":
                # Ответ компонента на одноразовый запрос выделения.
                # Важно: Streamlit компоненты возвращают "последнее значение" на каждом rerun,
                # поэтому обрабатываем selection только если он соответствует текущему запросу.
                req_uid = str(st.session_state.get("preview_request_uid") or "")
                evt_uid = str(evt_preview.get("request_uid") or "")
                if not req_uid or evt_uid != req_uid:
                    # Событие не относится к текущему запросу — игнорируем, чтобы не зациклить rerun.
                    pass
                else:
                    st.session_state["preview_request_selection"] = None
                    st.session_state["preview_request_uid"] = None

                    rng = evt_preview.get("range")
                    txt = (evt_preview.get("text") or "")
                    st.session_state["preview_selected_range"] = rng
                    st.session_state["preview_selected_text"] = txt

                    if st.session_state.get("preview_pending_action") == "rewrite":
                        st.session_state["preview_pending_action"] = None
                        sel_text = (txt or "").strip()
                        instr = _current_instr()
                        api_key = st.session_state.get("api_key") or os.getenv("OPENROUTER_API_KEY")

                        if not api_key:
                            st.error("API ключ не найден. Укажите OpenRouter API key слева (sk-or-v1-...).")
                        elif not rng or not rng.get("length"):
                            st.warning("Сначала выделите фрагмент в тексте.")
                        elif not sel_text:
                            st.warning("Пустое выделение.")
                        elif len(sel_text) > 4000:
                            st.error("Фрагмент слишком длинный (макс 4000 символов). Разбейте на части.")
                        elif not instr:
                            st.warning("Введите инструкцию для ИИ.")
                        else:
                            prompt = (
                                "Переделай фрагмент текста согласно инструкции.\n"
                                "Ответ дай только новым текстом без комментариев.\n\n"
                                f"Фрагмент:\n---\n{sel_text}\n---\n"
                                f"Инструкция: {instr}"
                            )
                            # Показываем спиннер справа от селектбокса инструкций,
                            # чтобы пользователь видел прогресс, даже если область предпросмотра вне экрана.
                            with cols_instr[1].spinner("ИИ переделывает выделение..."):
                                resp = generate_with_deepseek(api_key, prompt)
                            ai_text = None
                            if isinstance(resp, dict):
                                ai_text = resp.get("choices", [{}])[0].get("message", {}).get("content")
                            ai_text = (ai_text or "").strip()
                            if not ai_text:
                                st.error("Не удалось получить ответ от ИИ.")
                            else:
                                st.session_state["preview_rewrite_range"] = rng
                                st.session_state["preview_rewrite_source"] = sel_text
                                st.session_state["preview_rewrite_result"] = ai_text
                                # Нельзя менять ключи виджетов после их отрисовки в этом же прогоне.
                                # Поэтому кладём результат в буфер и делаем один rerun.
                                st.session_state["preview_fill_widgets"] = {"src": sel_text, "res": ai_text}

                                safe_rerun()

        # Редакторные блоки (Quill/streamlit-quill и их отладка) удалены.

        # Действия: сохранить, скачать .docx, очистить
        action_cols = st.columns([1, 1, 1])
        # Кнопка скрыта по просьбе — оставляем код, но не рендерим кнопку
        if False:
            title = (
                st.session_state.get("generated_title")
                or (st.session_state.get("gen_topic") or "План урока")
            )
            try:
                content_to_save = st.session_state.get("preview_html", "") or ""
                if not content_to_save:
                    src_md = st.session_state.get("generated_content") or st.session_state.get("stream_buffer") or ""
                    md = normalize_ai_markdown(_postprocess_plan_text(src_md))
                    content_to_save = quill_html_utils.sanitize_html_for_quill(markdown_to_html(md))
                plan = create_lesson_plan(
                    title=title,
                    subject=subject,
                    grade=grade,
                    topic=topic,
                    class_level=class_level,
                    content=content_to_save,
                    model_name=("local-template" if model_choice.startswith("Локальный") else "deepseek"),
                    visibility=visibility,
                    author_id=(getattr(current_user, "id", None) if current_user else None),
                )
                st.success(f"План сохранён (id={plan.id})")
            except Exception as e:
                st.error(f"Ошибка сохранения: {e}")

        docx_title = (
            st.session_state.get("generated_title")
            or (st.session_state.get("gen_topic") or "lesson_plan")
        )
        html_for_docx = st.session_state.get("preview_html", "") or ""
        if not html_for_docx:
            src_md = st.session_state.get("generated_content") or st.session_state.get("stream_buffer") or ""
            md = normalize_ai_markdown(_postprocess_plan_text(src_md))
            html_for_docx = quill_html_utils.sanitize_html_for_quill(markdown_to_html(md))
        docx_bytes = _html_to_docx_bytes(html_for_docx)
        try:
            action_cols[0].download_button(
                label="Скачать .docx",
                data=docx_bytes,
                file_name=_normalize_material_filename(
                    grade=st.session_state.get("gen_grade"),
                    kind="план",
                    topic=st.session_state.get("gen_topic") or docx_title,
                    ext="docx",
                ),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                on_click="ignore",
            )
        except KeyError:
            logging.exception("Missing media key when creating action download button")
            st.error("Файл временно недоступен. Пожалуйста, перезагрузите страницу и повторите попытку.")
        except Exception:
            logging.exception("Error while creating action download button")
            st.error("Не удалось подготовить файл для скачивания. Попробуйте ещё раз.")

        if action_cols[1].button("Очистить"):
            st.session_state["generated_content"] = ""
            st.session_state["stream_buffer"] = ""
            st.session_state["preview_html"] = ""
            st.session_state["preview_apply_replace"] = None
            safe_rerun()

        # Примечание: раньше тут был режим «вставьте выделенный текст». Теперь основная замена
        # делается прямо в редакторе через ПКМ (см. блок выше).
if st.session_state.get("show_recent_plans", False):
    with st.expander("Последние сохранённые планы уроков", expanded=False):
        plans = list_lesson_plans(limit=10)
        if plans:
            for plan in plans:
                meta = f"{plan.subject or 'Без предмета'} — {plan.grade or 'Без класса'}"
                if getattr(plan, "class_level", None):
                    meta += f" | уровень: {plan.class_level}"
                created = plan.created_at.strftime("%d.%m.%Y %H:%M") if plan.created_at else ""
                st.markdown(f"**{plan.title}**  ")
                st.caption(f"{meta} | создан: {created} | источник: {plan.model_name or 'не указан'}")
                with st.expander("Показать план"):
                    content = plan.content or ""
                    looks_like_html = content.lstrip().startswith("<")
                    if looks_like_html:
                        st.markdown(content, unsafe_allow_html=True)
                    else:
                        st.markdown(content)

                    download_ext = "html" if looks_like_html else "md"
                    download_mime = "text/html" if looks_like_html else "text/markdown"
                    try:
                        st.download_button(
                            label=f"Скачать как .{download_ext}",
                            data=content,
                            file_name=f"{_slugify(plan.title)}.{download_ext}",
                            mime=download_mime,
                            key=f"download_{plan.id}",
                            on_click="ignore",
                        )
                    except KeyError:
                        logging.exception("Missing media key when creating plan download button")
                        st.error("Файл временно недоступен. Пожалуйста, перезагрузите страницу и повторите попытку.")
                    except Exception:
                        logging.exception("Error while creating plan download button")
                        st.error("Не удалось подготовить файл для скачивания. Попробуйте ещё раз.")
        else:
            st.info("Пока нет сохранённых планов. Сгенерируйте первый план слева.")

if app_mode == "Пользовательский режим":
    st.header("Дополнительный материал для урока")

    tab_handout, tab_quiz, tab_test, tab_talk, tab_concept = st.tabs(
        ["Раздаточный материал", "Викторина", "Тест", "Беседа", "Конспект"]
    )

    # ----- Вкладка: Конспект (текущая работа с планом) -----
    with tab_concept:
        st.subheader("Конспект")
        
        # Показываем быстрый список последних конспектов
        with st.expander("Последние конспекты (файлы .docx)"):
            docs = sorted(materials_dir.glob("*.docx"))[:10]
            if docs:
                for p in docs:
                    st.write(p.name)
            else:
                st.write("Пока нет загруженных конспектов.")

    # ----- Вкладка: Раздаточный материал -----
    with tab_handout:
        st.subheader("Создать раздаточный материал")
        

        # Разметим форму в две равные колонки и распределим элементы поровну.
        col_left, col_right = st.columns(2)

        # Обработчик, который срабатывает при изменении входных параметров раздатки.
        def _handout_inputs_changed():
            st.session_state["handout_generated_md"] = ""
            st.session_state["handout_prompt_autofill"] = True

        # Синхронизируем предмет/класс/тему с верхними полями, чтобы не оставалось старых значений.
        gen_sub = st.session_state.get("gen_subject")
        gen_gr = st.session_state.get("gen_grade")
        gen_top = (st.session_state.get("gen_topic") or "").strip()

        if gen_sub:
            st.session_state["handout_subject"] = gen_sub
        else:
            st.session_state["handout_subject"] = st.session_state.get("handout_subject") or SUBJECTS[0]

        if gen_gr:
            st.session_state["handout_grade"] = gen_gr
        else:
            st.session_state["handout_grade"] = st.session_state.get("handout_grade") or GRADES[0]

        # Если тема изменилась — обновляем и сбрасываем ранее сгенерированный материал
        prev_handout_topic = (st.session_state.get("handout_topic") or "").strip()
        if gen_top != prev_handout_topic:
            st.session_state["handout_topic"] = gen_top
            st.session_state["handout_generated_md"] = ""
            st.session_state["handout_prompt_autofill"] = True

        handout_subject = st.session_state.get("handout_subject")
        handout_grade = st.session_state.get("handout_grade")
        handout_topic = st.session_state.get("handout_topic")

        # Левый столбец: краткая информация, формат и длительность, дополнительные элементы
        with col_left:
            st.markdown(
                f"<div style='font-size:18px;color:#000;font-weight:600'>Используем: {html.escape(str(handout_subject))} • {html.escape(str(handout_grade))} • Тема: {html.escape(str(handout_topic or '—'))}</div>",
                unsafe_allow_html=True,
            )

            handout_kind = st.selectbox(
                "Формат материала",
                ["карточки-задания", "рабочий лист", "опорный конспект", "таблица", "схема", "квест"],
                index=1,
                key="handout_kind",
                on_change=_handout_inputs_changed,
            )

            handout_time_min = st.slider(
                "Длительность (мин)",
                min_value=5,
                max_value=40,
                value=int(st.session_state.get("handout_time_min", 15) or 15),
                step=5,
                key="handout_time_min",
                on_change=_handout_inputs_changed,
            )

            handout_elements = st.multiselect(
                "Задания должны включать",
                [
                    "теорию в рамочках",
                    "примеры",
                    "проблемные вопросы",
                    "тесты с выбором ответа",
                    "задания на анализ",
                    "творческие задания",
                    "рефлексия",
                ],
                default=[],
                key="handout_elements",
                on_change=_handout_inputs_changed,
                placeholder="выберите дополнительные параметры",
            )

        # Правый столбец: режим работы, сложность и поле заметок
        with col_right:
            handout_work_mode = st.selectbox(
                "Формат работы",
                ["индивидуальная работа", "парная работа", "групповая работа"],
                index=0,
                key="handout_work_mode",
                on_change=_handout_inputs_changed,
            )

            handout_difficulty = st.selectbox(
                "Уровень сложности",
                ["★ Базовый", "★★ Средний", "★★★ Повышенный"],
                index=1,
                key="handout_difficulty",
                on_change=_handout_inputs_changed,
            )

            handout_notes = st.text_area("Примечания (опционально)", key="handout_notes", on_change=_handout_inputs_changed)

        if "handout_prompt" not in st.session_state:
            st.session_state["handout_prompt"] = ""

        # Streamlit-нюанс: если у виджета задан key, параметр `value=`
        # применяется только при первом рендере. Поэтому держим флаг,
        # можно ли автообновлять промпт при смене селектов.
        if "handout_prompt_autofill" not in st.session_state:
            st.session_state["handout_prompt_autofill"] = True

        def _handout_prompt_mark_dirty() -> None:
            st.session_state["handout_prompt_autofill"] = False

        if "handout_generated_md" not in st.session_state:
            st.session_state["handout_generated_md"] = ""

        auto_prompt = (
            "Ты — опытный учитель-методист. Сгенерируй раздаточный материал ДЛЯ УЧЕНИКА по заданным параметрам.\n"
            "Вывод: строго в Markdown (без HTML), чтобы можно было сохранить в .docx.\n\n"
            f"Предмет: {handout_subject}.\n"
            f"Класс: {handout_grade}.\n"
            f"Тема: {handout_topic}.\n\n"
            f"Формат материала: {handout_kind}.\n"
            f"Формат работы: {handout_work_mode}.\n"
            f"Длительность: ~{int(st.session_state.get('handout_time_min') or 15)} минут.\n"
            f"Уровень сложности: {handout_difficulty}.\n\n"
            "Общие требования:\n"
            "- Структура: Заголовок → Краткая цель (1–2 строки) → Теория (если нужна) → Задания → Самопроверка (по желанию).\n"
            "- Язык и объём: соответствуй классу и длительности; без воды.\n"
            "- Форматирование: используй списки, таблицы (если уместно), разделители.\n"
            "- Для 'карточки-задания' сделай 4–8 карточек с короткими заданиями.\n"
            "- Для 'рабочий лист' сделай последовательный лист с полями для ответа (подчёркивания/пустые строки).\n"
            "- Для 'опорный конспект' сделай краткие пункты + ключевые определения/формулы + 1–2 примера.\n"
            "- Для 'таблица' дай таблицу для заполнения + 3–6 вопросов к ней.\n"
            "- Для 'схема' дай схему в виде иерархического списка/блоков + вопросы.\n"
            "- Для 'квест' сделай 5–8 шагов с подсказками и итоговым заданием.\n"
            "- Если есть тесты с выбором ответа: 4 варианта (A–D) и в конце добавь небольшой 'Ключ ответов'.\n"
            "- Не пиши рассуждений, только готовый материал.\n"
        )
        if handout_elements:
            auto_prompt += "\nЗадания обязательно должны включать: " + ", ".join(handout_elements) + ".\n"
        if (handout_notes or "").strip():
            auto_prompt += f"\nДополнительные пожелания учителя:\n{handout_notes.strip()}\n"

        # Если пользователь не правил промпт вручную, синхронизируем редактор с авто-промптом.
        if st.session_state.get("handout_prompt_autofill"):
            st.session_state["handout_prompt_editor"] = auto_prompt

        if False:
            with st.expander("Шаблон для генерации (можно править)", expanded=False):
                st.caption("Можно оставить как есть — он собран автоматически из формы выше.")
                prompt_cols = st.columns([1, 1, 3])
                if prompt_cols[0].button("Сбросить к авто", key="handout_prompt_reset_btn"):
                    st.session_state["handout_prompt_autofill"] = True
                    st.session_state["handout_prompt_editor"] = auto_prompt

                st.session_state["handout_prompt"] = st.text_area(
                    "Промпт",
                    height=220,
                    key="handout_prompt_editor",
                    on_change=_handout_prompt_mark_dirty,
                )

        gen_cols = st.columns([1, 1, 2])
        if gen_cols[0].button("Сгенерировать раздаточный материал (ИИ)", key="handout_generate_btn"):
            api_key_local = st.session_state.get("api_key") or os.getenv("OPENROUTER_API_KEY")
            if not api_key_local:
                st.error("Укажите OpenRouter API key в админ-панели (sk-or-v1-...).")
            elif not (handout_topic or "").strip():
                st.warning("Заполните тему раздаточного материала.")
            else:
                prompt_text = (st.session_state.get("handout_prompt_editor") or auto_prompt).strip()
                with st.spinner("Генерирую раздатку..."):
                    resp = generate_with_deepseek(api_key_local, prompt_text)
                md_text = (
                    resp.get("choices", [{}])[0]
                    .get("message", {})
                    .get("content", "")
                )
                st.session_state["handout_generated_md"] = (md_text or "").strip()

        if gen_cols[1].button("Очистить", key="handout_clear_btn"):
            st.session_state["handout_generated_md"] = ""

        if (st.session_state.get("handout_generated_md") or "").strip():
            st.markdown("### Результат (можно править)")
            st.session_state["handout_generated_md"] = st.text_area(
                "handout_generated_md_editor",
                value=st.session_state["handout_generated_md"],
                height=320,
                label_visibility="collapsed",
            )

            save_cols = st.columns([1, 1, 2])
            handout_docx_title = _normalize_material_filename(
                grade=handout_grade,
                kind="раздатка",
                topic=handout_topic or "материал",
                ext="docx",
            )

            # Показываем одну кнопку-скачивалку, которая отдаёт .docx напрямую (без записи на диск)
            try:
                src_md = st.session_state.get("handout_generated_md") or ""
                md_norm = normalize_ai_markdown(_postprocess_plan_text(src_md))
                html_for_docx = quill_html_utils.sanitize_html_for_quill(markdown_to_html(md_norm))
                bytes_docx = _html_to_docx_bytes(html_for_docx)

                from storage import create_material

                topics_csv = (handout_topic or "").strip() or None
                user_id = st.session_state.get("user_id")

                def _register_handout():
                    try:
                        create_material(
                            filename=handout_docx_title,
                            uploader_id=user_id,
                            topics=topics_csv,
                            path=None,
                            subject=handout_subject,
                            grade=handout_grade,
                        )
                        st.session_state["handout_last_registered"] = True
                    except Exception:
                        logging.exception("Error while registering handout in DB")

                # Отдаём файл клиенту напрямую — кнопка инициирует скачивание в браузере.
                try:
                    save_cols[0].download_button(
                        "Скачать раздаточный материал",
                        data=bytes_docx,
                        file_name=handout_docx_title,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        on_click=_register_handout,
                    )
                except Exception:
                    logging.exception("Error while creating download button for handout")
                    st.error("Не удалось подготовить файл для скачивания. Попробуйте ещё раз.")

            except Exception as e:
                logging.exception("Error while preparing handout bytes for download")
                st.error(f"Ошибка при подготовке раздатки: {e}")

        # Блок "Создать раздатку из текущего плана" временно скрыт
        # (удален по запросу пользователя — можно вернуть при необходимости)

    # ----- Вкладка: Викторина -----
    with tab_quiz:
        st.subheader("Викторина (ИИ)")

        def _quiz_mark_prompt_dirty():
            st.session_state["quiz_prompt_autofill"] = False

        def _quiz_params_changed():
            st.session_state["quiz_generated_json_text"] = ""

        # Берём тему/предмет/класс из раздатки, если пользователь менял их там;
        # иначе — из формы генерации плана. Так меньше путаницы "почему не по теме".
        quiz_subject = (
            st.session_state.get("handout_subject")
            or st.session_state.get("gen_subject")
            or SUBJECTS[0]
        )
        quiz_grade = (
            st.session_state.get("handout_grade")
            or st.session_state.get("gen_grade")
            or GRADES[0]
        )
        quiz_topic = (
            (st.session_state.get("handout_topic") or "").strip()
            or (st.session_state.get("gen_topic") or "").strip()
        )

        # Авто-обновление названия викторины при изменении темы,
        # если пользователь не редактировал поле `quiz_title` вручную.
        if "quiz_title_autofill" not in st.session_state:
            st.session_state["quiz_title_autofill"] = True
        # Если автозаполнение включено — синхронизируем значение в session_state
        if st.session_state.get("quiz_title_autofill"):
            st.session_state["quiz_title"] = quiz_topic or "Викторина"

        st.markdown(
            f"<div style='font-size:18px;color:#000;font-weight:600'>Используем для викторины: {html.escape(str(quiz_subject))} • {html.escape(str(quiz_grade))} • Тема: {html.escape(str(quiz_topic or '—'))}</div>",
            unsafe_allow_html=True,
        )

        params_cols = st.columns([2, 1, 2])
        # Название викторины формируется автоматически из темы и хранится в session_state.
        quiz_title = st.session_state.get("quiz_title", (quiz_topic or "Викторина"))
        quiz_delivery = params_cols[0].selectbox(
            "Формат проведения",
            [
                "Устный (ведущий читает — дети отвечают)",
                "Письменный (распечатанные бланки / рабочие листы)",
                "Интерактивный (презентация, кликер)",
                "Смешанный (устно + письменно)",
            ],
            index=0,
            key="quiz_delivery",
            on_change=_quiz_params_changed,
        )
        quiz_types = params_cols[0].multiselect(
            "Форматы",
            [
                "Выбор одного ответа",
                "Несколько ответов",
                "Верно/неверно",
                "Короткий ответ",
            ],
            default=["Выбор одного ответа", "Верно/неверно"],
            key="quiz_types",
            on_change=_quiz_params_changed,
        )
        quiz_num_questions = params_cols[1].slider(
            "Вопросов",
            min_value=5,
            max_value=30,
            value=10,
            step=1,
            key="quiz_num_questions",
            on_change=_quiz_params_changed,
        )
        quiz_difficulty = params_cols[2].selectbox(
            "Сложность",
            ["Смешанная", "Лёгкая", "Средняя", "Сложная"],
            index=0,
            key="quiz_difficulty",
            on_change=_quiz_params_changed,
        )

        quiz_mode = params_cols[2].selectbox(
            "Тип викторины",
            ["Индивидуальная", "Командная", "Парная", "Общеклассная"],
            index=0,
            key="quiz_mode",
            on_change=_quiz_params_changed,
        )

        quiz_dynamics = params_cols[2].selectbox(
            "Динамика",
            [
                "Соревновательная (баллы, очки, жетоны)",
                "Без соревнования (познавательная)",
                "С призами / наклейками / поощрениями",
            ],
            index=1,
            key="quiz_dynamics",
            on_change=_quiz_params_changed,
        )

        quiz_notes = params_cols[0].text_area(
            "Примечания (опционально)",
            value=st.session_state.get("quiz_notes", ""),
            key="quiz_notes",
            on_change=_quiz_params_changed,
            height=80,
        )
        quiz_time_min = params_cols[1].slider(
            "Ориентировочное время (мин)",
            min_value=5,
            max_value=40,
            value=int(st.session_state.get("quiz_time_min", 15) or 15),
            step=5,
            key="quiz_time_min",
            on_change=_quiz_params_changed,
        )

        # Авто-промпт для викторины
        types_hint = ", ".join(quiz_types) if quiz_types else "на твой выбор"
        auto_prompt = (
            "Ты — опытный учитель и методист. Составь школьную викторину для проверки понимания темы.\n\n"
            "ВАЖНО: все вопросы должны быть строго по теме, указанной ниже.\n\n"
            f"Предмет: {quiz_subject}.\n"
            f"Класс: {quiz_grade}.\n"
            "ТЕМА (главная):\n"
            f"\"\"\"\n{quiz_topic or '(тема не указана)'}\n\"\"\"\n"
            f"Название викторины: {quiz_title or 'Викторина'}.\n"
            f"Количество вопросов: {int(quiz_num_questions)}.\n"
            f"Сложность: {quiz_difficulty}.\n"
            f"Тип викторины: {st.session_state.get('quiz_mode', 'Индивидуальная')}.\n"
            f"Формат проведения: {st.session_state.get('quiz_delivery', 'Устный (ведущий читает — дети отвечают)')}.\n"
            f"Динамика: {st.session_state.get('quiz_dynamics', 'Без соревнования (познавательная)')}.\n"
            f"Форматы вопросов: {types_hint}.\n"
            f"Время выполнения: ~{int(quiz_time_min)} минут.\n\n"
            "Требования к качеству:\n"
            "- Уровень и язык: строго по возрасту/классу, без спорных/взрослых тем.\n"
            "- Проверяй факты, избегай двусмысленности формулировок.\n"
            "- Если тема содержит математику/формулы: НЕ используй LaTeX и обратные слэши; пиши формулы обычным текстом.\n"
            "- Для вопросов с вариантами ответов: ровно 4 варианта (A–D), один или несколько верных по типу.\n"
            "- Для 'Верно/неверно': варианты ['Верно','Неверно'].\n"
            "- Для 'Короткий ответ': дай краткий ожидаемый ответ (строка).\n"
            "- Добавь краткое пояснение (1 предложение) почему ответ верный — в поле explanation.\n\n"
            "Верни СТРОГО валидный JSON (без markdown, без ```), в следующей структуре:\n"
            "{\n"
            "  'title': string,\n"
            "  'subject': string,\n"
            "  'grade': string,\n"
            "  'topic': string,\n"
            "  'difficulty': string,\n"
            "  'time_minutes': number,\n"
            "  'questions': [\n"
            "    {\n"
            "      'id': number,\n"
            "      'type': 'single_choice'|'multiple_choice'|'true_false'|'short_answer',\n"
            "      'question': string,\n"
            "      'choices': string[],\n"
            "      'answer': number|string|number[],\n"
            "      'explanation': string,\n"
            "      'points': number\n"
            "    }\n"
            "  ]\n"
            "}\n\n"
            "Правила для answer:\n"
            "- single_choice/true_false: индекс правильного варианта (0..3 или 0..1 для true_false)\n"
            "- multiple_choice: список индексов (например [0,2])\n"
            "- short_answer: строка\n"
        )

        st.session_state.setdefault("quiz_prompt_autofill", True)
        if st.session_state.get("quiz_prompt_autofill"):
            st.session_state["quiz_prompt_editor"] = auto_prompt
        else:
            st.info("Промпт был изменён вручную и не автообновляется. Нажмите «Сбросить к авто», если тема/параметры изменились.")

        if False:
            with st.expander("Шаблон для генерации (можно править)", expanded=False):
                prompt_cols = st.columns([1, 3])
                if prompt_cols[0].button("Сбросить к авто", key="quiz_prompt_reset_btn"):
                    st.session_state["quiz_prompt_autofill"] = True
                    st.session_state["quiz_prompt_editor"] = auto_prompt

                st.text_area(
                    "Промпт",
                    height=240,
                    key="quiz_prompt_editor",
                    on_change=_quiz_mark_prompt_dirty,
                )

        gen_cols = st.columns([1, 1, 2])
        if gen_cols[0].button("Сгенерировать викторину (ИИ)", key="quiz_generate_btn"):
            api_key_local = st.session_state.get("api_key") or os.getenv("OPENROUTER_API_KEY")
            if not api_key_local:
                st.error("Укажите OpenRouter API key в админ-панели (sk-or-v1-...).")
            elif not quiz_topic:
                st.warning("Сначала задайте тему (в форме плана урока).")
            else:
                prompt_text = (st.session_state.get("quiz_prompt_editor") or auto_prompt).strip()
                with st.spinner("Генерирую викторину..."):
                    resp = generate_with_deepseek(api_key_local, prompt_text)
                raw = (
                    resp.get("choices", [{}])[0]
                    .get("message", {})
                    .get("content", "")
                )

                def _extract_json(text: str) -> str:
                    if not text:
                        return ""
                    t = text.strip()
                    if t.startswith("```"):
                        t = re.sub(r"^```[a-zA-Z0-9_-]*\s*", "", t)
                        t = re.sub(r"```\s*$", "", t).strip()
                    m = re.search(r"\{.*\}", t, flags=re.S)
                    return (m.group(0) if m else t).strip()

                extracted = _extract_json(raw)
                # Попытка получить валидный JSON. Если AI вернул python-словарь
                # с одинарными кавычками, пробуем ast.literal_eval как запасной вариант.
                pretty = extracted
                try:
                    parsed = json.loads(extracted)
                    pretty = json.dumps(parsed, ensure_ascii=False, indent=2)
                except Exception:
                    try:
                        import ast

                        parsed = ast.literal_eval(extracted)
                        # Конвертируем в корректный JSON-строковый формат
                        pretty = json.dumps(parsed, ensure_ascii=False, indent=2)
                    except Exception:
                        # Оставляем оригинальный извлечённый текст для ручной правки
                        pretty = extracted

                st.session_state["quiz_generated_json_text"] = pretty

        if gen_cols[1].button("Очистить", key="quiz_clear_btn"):
            st.session_state["quiz_generated_json_text"] = ""

        def _quiz_to_markdown(q: dict) -> str:
            title = (q.get("title") or quiz_title or "Викторина").strip()
            subject = (q.get("subject") or quiz_subject or "").strip()
            grade = (q.get("grade") or quiz_grade or "").strip()
            topic = (q.get("topic") or quiz_topic or "").strip()
            difficulty = (q.get("difficulty") or quiz_difficulty or "").strip()
            time_minutes = q.get("time_minutes") or quiz_time_min

            lines = [f"# {title}"]
            meta = []
            if subject:
                meta.append(f"Предмет: {subject}")
            if grade:
                meta.append(f"Класс: {grade}")
            if topic:
                meta.append(f"Тема: {topic}")
            if difficulty:
                meta.append(f"Сложность: {difficulty}")
            if time_minutes:
                meta.append(f"Время: ~{int(time_minutes)} мин")
            if meta:
                lines.append("\n" + "  ".join(meta) + "\n")

            questions = q.get("questions") or []
            answer_key = []
            type_labels = {
                "single_choice": "Выбор одного ответа",
                "multiple_choice": "Несколько ответов",
                "true_false": "Верно/неверно",
                "short_answer": "Короткий ответ",
            }

            for idx, qq in enumerate(questions, start=1):
                q_type = (qq.get("type") or "").strip()
                q_text = (qq.get("question") or "").strip()
                choices = qq.get("choices") or []
                answer = qq.get("answer")
                explanation = (qq.get("explanation") or "").strip()

                lines.append(f"## Вопрос {idx}")
                if q_type:
                    lines.append(f"Тип: {type_labels.get(q_type, q_type)}")
                lines.append(q_text or "(без текста вопроса)")

                if choices:
                    letters = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
                    for c_i, c in enumerate(choices):
                        prefix = letters[c_i] if c_i < len(letters) else str(c_i + 1)
                        lines.append(f"- {prefix}. {(str(c) or '').strip()}")

                answer_str = ""
                if isinstance(answer, list):
                    try:
                        answer_str = ", ".join(str(a) for a in answer)
                    except Exception:
                        answer_str = str(answer)
                else:
                    answer_str = "" if answer is None else str(answer)

                if answer_str or explanation:
                    key_line = f"{idx}. Ответ: {answer_str}" if answer_str else f"{idx}. Ответ: (не указан)"
                    if explanation:
                        key_line += f" — {explanation}"
                    answer_key.append(key_line)

                lines.append("")

            if answer_key:
                lines.append("---")
                lines.append("# Ключ ответов")
                # Добавляем ключ ответов как сырые HTML-параграфы, чтобы
                # при конвертации Markdown->HTML->DOCX они не были объединены
                # в продолжающийся нумерованный список. Это гарантирует,
                # что нумерация начнётся заново в разделе ключа.
                for ak in answer_key:
                    lines.append(f"<p>{ak}</p>")

            return "\n".join(lines).strip() + "\n"

        # Результат
        if (st.session_state.get("quiz_generated_json_text") or "").strip():
            quiz_text = st.session_state["quiz_generated_json_text"]

            quiz_obj = None
            parse_error = None
            try:
                quiz_obj = json.loads(quiz_text)
            except Exception as e:
                parse_error = e

            if quiz_obj and isinstance(quiz_obj, dict):
                st.markdown("### Результат")
                # Показываем человекочитаемый вид викторины в компактном окне.
                try:
                    quiz_md_preview = _quiz_to_markdown(quiz_obj)
                    quiz_html = markdown_to_html(quiz_md_preview)
                    # Увеличиваем размер именно для заголовков с конкретными словами
                    try:
                        from bs4 import BeautifulSoup

                        soup = BeautifulSoup(quiz_html, "html.parser")
                        # сравниваем без учёта регистра, ищем фразы внутри текста заголовка
                        special = {"раздаточный материал", "конспект", "викторина"}
                        for tag in soup.find_all(["h1", "h2", "h3"]):
                            text = (tag.get_text(strip=True) or "").lower()
                            if any(text == s or s in text for s in special):
                                # Добавляем/обновляем inline-стиль для явного увеличения шрифта заголовка
                                existing = (tag.get("style", "") or "").strip()
                                existing = existing.rstrip(";")
                                new_styles = "font-size:2.0rem; font-weight:800;"
                                tag["style"] = (f"{existing}; {new_styles}" if existing else new_styles)
                        quiz_html = str(soup)
                    except Exception:
                        # если BeautifulSoup недоступен или что-то пошло не так, продолжаем с исходным HTML
                        pass

                    preview_html = (
                        "<style>"
                        ".qa-preview h1{font-size:1.25rem;margin:0.4rem 0;}"
                        ".qa-preview h2{font-size:1.1rem;margin:0.35rem 0;}"
                        ".qa-preview h3{font-size:1rem;margin:0.3rem 0;}"
                        ".qa-preview p{margin:0.25rem 0;}"
                        "</style>"
                        f"<div class='qa-preview' style='max-height:260px; overflow:auto; padding:0.6rem; border:1px solid #e6e6e6; background:#ffffff; border-radius:6px'>{quiz_html}</div>"
                    )
                    st.markdown(preview_html, unsafe_allow_html=True)
                except Exception:
                    st.warning("Не удалось сформировать превью викторины. Откройте JSON в экспандере для отладки.")

                # Экспандер для просмотра/правки JSON скрыт по умолчанию (отладка удалена).
            else:
                # Некорректный JSON — показываем редактор сразу и предупреждение
                st.markdown("### Результат (JSON — ошибка парсинга)")
                st.warning(f"Не удалось распарсить JSON: {parse_error}")
                st.session_state["quiz_generated_json_text"] = st.text_area(
                    "quiz_generated_json_editor",
                    value=st.session_state["quiz_generated_json_text"],
                    height=320,
                    label_visibility="collapsed",
                )

            

            if isinstance(quiz_obj, dict):
                try:
                    quiz_md = _quiz_to_markdown(quiz_obj)
                    md_norm = normalize_ai_markdown(_postprocess_plan_text(quiz_md))
                    html_for_docx = quill_html_utils.sanitize_html_for_quill(markdown_to_html(md_norm))
                    bytes_docx = _html_to_docx_bytes(html_for_docx)

                    st.download_button(
                        "Скачать викторину (.docx)",
                        data=bytes_docx,
                        file_name=_normalize_material_filename(
                            grade=quiz_grade,
                            kind="викторина",
                            topic=quiz_topic or quiz_title or "викторина",
                            ext="docx",
                        ),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        on_click="ignore",
                        key="quiz_download_docx_btn",
                    )
                except Exception as e:
                    logging.exception("Error while converting quiz to docx")
                    st.error(f"Не удалось собрать .docx: {e}")

        # Пустой шаблон для ручного заполнения удалён по просьбе пользователя.
        
    # ----- Вкладка: Тест -----
    with tab_test:
        st.subheader("Тест (сбор параметров)")

        # Берём тему/предмет/класс из раздатки или формы плана
        test_subject = (
            st.session_state.get("handout_subject")
            or st.session_state.get("gen_subject")
            or SUBJECTS[0]
        )
        test_grade = (
            st.session_state.get("handout_grade")
            or st.session_state.get("gen_grade")
            or GRADES[0]
        )
        test_topic = (
            (st.session_state.get("handout_topic") or "").strip()
            or (st.session_state.get("gen_topic") or "").strip()
        )

        st.markdown(
            f"<div style='font-size:18px;color:#000;font-weight:600'>Используем для теста: {html.escape(str(test_subject))} • {html.escape(str(test_grade))} • Тема: {html.escape(str(test_topic or '—'))}</div>",
            unsafe_allow_html=True,
        )

        def _compact_spinner(container, label: str, key: str, default: int = 2, max_value: int = 100, on_change=None) -> int:
            left_col, right_col = container.columns([6, 1])
            left_col.markdown(label)
            # If the key already exists in session_state, avoid passing an explicit
            # `value` to the widget — Streamlit will initialize the widget from
            # session_state. Passing both causes a warning about default vs session state.
            if key in st.session_state:
                value = right_col.number_input(
                    label,
                    min_value=0,
                    max_value=max_value,
                    step=1,
                    format="%d",
                    key=key,
                    label_visibility="collapsed",
                    on_change=on_change,
                )
            else:
                value = right_col.number_input(
                    label,
                    min_value=0,
                    max_value=max_value,
                    value=int(default),
                    step=1,
                    format="%d",
                    key=key,
                    label_visibility="collapsed",
                    on_change=on_change,
                )
            return int(value)

        info_col1, info_col2, info_col3 = st.columns(3)

        with info_col1:
            st.markdown(
                "<div style='margin-top:0.2rem;margin-bottom:0.1rem;font-weight:800'>Цель теста:</div>",
                unsafe_allow_html=True,
            )
            test_goal = st.selectbox(
                "Цель теста",
                [
                    "Текущая проверка",
                    "Итоговый контроль",
                    "Входная диагностика",
                    "Самопроверка",
                    "Подготовка к контрольной работе",
                ],
                index=0,
                key="test_goal",
                label_visibility="collapsed",
            )

            st.markdown("**Типы вопросов (закрытые):**")
            test_single = _compact_spinner(info_col1, "Выбор одного ответа (1 правильный, 3-4 варианта)", "test_single")
            test_multiple = _compact_spinner(info_col1, "Выбор нескольких ответов", "test_multiple")
            test_truefalse = _compact_spinner(info_col1, "Верно/Неверно", "test_truefalse")
            test_find_extra = _compact_spinner(info_col1, "Найди лишнее", "test_find_extra")

        with info_col2:
            st.markdown("**Типы вопросов (открытые):**")
            test_match = _compact_spinner(info_col2, "На соответствие (соединить колонки)", "test_match", default=1)
            test_sequence = _compact_spinner(info_col2, "На последовательность (расставить по порядку)", "test_sequence", default=1)
            test_short = _compact_spinner(info_col2, "Краткий ответ (слово/словосочетание)", "test_short", default=1)
            test_long = _compact_spinner(info_col2, "Развёрнутый ответ (1-2 предложения)", "test_long", default=1)
            # Переносим блок примечаний в средний столбец
            st.markdown("<div style='font-weight:600;margin-bottom:0.2rem'>Доп. примечания (опционально):</div>", unsafe_allow_html=True)
            test_notes = st.text_area("Доп. примечания (опционально)", value="", key="test_notes", height=130, label_visibility="collapsed")

            # Наглядные вопросы удалены по запросу пользователя
            test_by_image = 0
            test_by_table = 0

        # Авто-распределение по сложности в зависимости от общего числа вопросов
        def _mark_diff_manual():
            st.session_state["test_diff_manual"] = True

        # Собираем общее количество вопросов из полей выше
        total_questions = sum(
            int(x or 0)
            for x in (
                st.session_state.get("test_single"),
                st.session_state.get("test_multiple"),
                st.session_state.get("test_truefalse"),
                st.session_state.get("test_find_extra"),
                st.session_state.get("test_match"),
                st.session_state.get("test_sequence"),
                st.session_state.get("test_short"),
                st.session_state.get("test_long"),
            )
        )

        # Инициализация флага автозаполнения
        if "test_diff_manual" not in st.session_state:
            st.session_state["test_diff_manual"] = False

        # Если пользователь не редактировал вручную — установим дефолтное распределение
        if not st.session_state.get("test_diff_manual", False):
            # Процентное распределение: базовый 60%, повышенный 30%, высокий 10% (в сумме = total)
            basic = int(round(total_questions * 0.6))
            mid = int(round(total_questions * 0.3))
            high = max(0, total_questions - basic - mid)
            st.session_state["test_diff_basic"] = basic
            st.session_state["test_diff_mid"] = mid
            st.session_state["test_diff_high"] = high

        with info_col3:
            st.markdown("**Распределение по сложности:**")
            diff_basic = _compact_spinner(
                info_col3,
                "Базовый уровень",
                "test_diff_basic",
                default=st.session_state.get("test_diff_basic", 2),
                max_value=200,
                on_change=_mark_diff_manual,
            )
            diff_mid = _compact_spinner(
                info_col3,
                "Повышенный уровень",
                "test_diff_mid",
                default=st.session_state.get("test_diff_mid", 2),
                max_value=200,
                on_change=_mark_diff_manual,
            )
            diff_high = _compact_spinner(
                info_col3,
                "Высокий уровень",
                "test_diff_high",
                default=st.session_state.get("test_diff_high", 2),
                max_value=200,
                on_change=_mark_diff_manual,
            )

            st.markdown("**Дополнительные пожелания:**")
            add_keys = st.checkbox("Добавить ключи с ответами (отдельно)", value=True, key="test_add_keys")
            add_comments = st.checkbox("Добавить комментарии к сложным вопросам", value=False, key="test_add_comments")
            add_variants = st.checkbox("Сделать 2 варианта (одинаковой сложности)", value=False, key="test_add_variants")
            add_local = st.checkbox("Учесть местный / региональный контекст (Беларусь)", value=False, key="test_add_local")

            # Примечания перенесены в средний столбец

        gen_cols = st.columns([1, 1, 2])
        if gen_cols[0].button("Сгенерировать тест (ИИ)", key="test_generate_btn"):
            api_key_local = st.session_state.get("api_key") or os.getenv("OPENROUTER_API_KEY")
            if not api_key_local:
                st.error("Укажите OpenRouter API key в админ-панели (sk-or-v1-...).")
            elif not test_topic:
                st.warning("Сначала задайте тему (в форме плана урока).")
            else:
                # Собираем промпт по шаблону пользователя
                prompt_lines = []
                prompt_lines.append("Ты — экзаменатор/методист. Составь тест по заданной теме.")
                prompt_lines.append(f"Предмет: {test_subject}.")
                prompt_lines.append(f"Класс: {test_grade}.")
                prompt_lines.append("ТЕМА:")
                prompt_lines.append(f'"""\n{test_topic}\n"""')
                # test_goal — единый выбор
                prompt_lines.append("Цель теста: " + (test_goal or "—"))
                prompt_lines.append("Типы вопросов и количества:")
                prompt_lines.append(f"- Выбор одного ответа: {int(test_single)}")
                prompt_lines.append(f"- Выбор нескольких ответов: {int(test_multiple)}")
                prompt_lines.append(f"- Верно/Неверно: {int(test_truefalse)}")
                prompt_lines.append(f"- Найди лишнее: {int(test_find_extra)}")
                prompt_lines.append(f"- На соответствие: {int(test_match)}")
                prompt_lines.append(f"- На последовательность: {int(test_sequence)}")
                prompt_lines.append(f"- Краткий ответ: {int(test_short)}")
                prompt_lines.append(f"- Развёрнутый ответ: {int(test_long)}")
                # Наглядные вопросы опущены
                prompt_lines.append("Распределение по сложности:")
                prompt_lines.append(f"- Базовый: {int(diff_basic)}; Повышенный: {int(diff_mid)}; Высокий: {int(diff_high)}")
                prompt_lines.append("Доп. пожелания:")
                if add_keys:
                    prompt_lines.append("- Добавить ключ ответов отдельно")
                if add_comments:
                    prompt_lines.append("- Добавить комментарии к сложным вопросам")
                if add_variants:
                    prompt_lines.append("- Сделать 2 варианта одинаковой сложности")
                if add_local:
                    prompt_lines.append("- Учитывать местный контекст: Беларусь")
                if test_notes and test_notes.strip():
                    prompt_lines.append(f"Примечания учителя: {test_notes.strip()}")

                prompt_text = "\n".join(prompt_lines)
                with st.spinner("Генерирую тест..."):
                    resp = generate_with_deepseek(api_key_local, prompt_text)
                text = (
                    resp.get("choices", [{}])[0]
                    .get("message", {})
                    .get("content", "")
                )
                st.session_state["test_generated_md"] = (text or "").strip()

        if gen_cols[1].button("Очистить", key="test_clear_btn"):
            st.session_state["test_generated_md"] = ""

        if st.session_state.get("test_generated_md"):
            st.markdown("### Результат (можно править)")
            st.session_state["test_generated_md"] = st.text_area(
                "test_generated_md_editor",
                value=st.session_state.get("test_generated_md", ""),
                height=420,
                label_visibility="collapsed",
            )

            try:
                test_docx_title = _normalize_material_filename(
                    grade=test_grade,
                    kind="тест",
                    topic=test_topic or "тест",
                    ext="docx",
                )
                src_md = st.session_state.get("test_generated_md") or ""
                md_norm = normalize_ai_markdown(_postprocess_plan_text(src_md))
                html_for_docx = quill_html_utils.sanitize_html_for_quill(markdown_to_html(md_norm))
                bytes_docx = _html_to_docx_bytes(html_for_docx)

                st.download_button(
                    "Скачать тест (.docx)",
                    data=bytes_docx,
                    file_name=test_docx_title,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    on_click="ignore",
                    key="test_download_docx_btn",
                )
            except Exception as e:
                logging.exception("Error while converting test to docx")
                st.error(f"Не удалось собрать .docx: {e}")

        # ----- Вкладка: Беседа -----
        with tab_talk:
            st.subheader("Беседа (ИИ)")

            def _talk_params_changed():
                st.session_state["talk_generated_md"] = ""
                st.session_state["talk_prompt_autofill"] = True

            # Берём тему и класс из верхней формы плана урока
            talk_subject = st.session_state.get("gen_subject") or st.session_state.get("handout_subject") or SUBJECTS[0]
            talk_grade = st.session_state.get("gen_grade") or st.session_state.get("handout_grade") or GRADES[0]
            talk_topic = (st.session_state.get("gen_topic") or "").strip()

            st.markdown(
                f"<div style='font-size:18px;color:#000;font-weight:600'>Используем: {html.escape(str(talk_subject))} • {html.escape(str(talk_grade))} • Тема: {html.escape(str(talk_topic or '—'))}</div>",
                unsafe_allow_html=True,
            )

            left, right = st.columns(2)

            with left:
                talk_goal = st.selectbox(
                    "Цель и педагогический фокус",
                    [
                        "—",
                        "Знакомство с темой (первичное информирование).",
                        "Воспитательный момент (формирование ценностей: доброта, уважение к истории, экологичность).",
                        "Интерактив и вовлечение (чтобы дети больше говорили и делали).",
                        "Эмоциональный отклик (вызвать удивление, радость, сопереживание).",
                        "Практический навык (научиться чему-то простому за время беседы).",
                        "Подготовка к проекту (беседа как вводная часть для дальнейшей работы).",
                    ],
                    index=0,
                    key="talk_goal",
                    on_change=_talk_params_changed,
                )

                talk_format = st.selectbox(
                    "Формат и методы",
                    [
                        "—",
                        "Монолог-рассказ учителя с элементами вопросов.",
                        "Эвристическая беседа (учитель задает наводящие вопросы, дети сами приходят к выводам).",
                        "Беседа с опорой на наглядность (укажите, если у вас есть конкретные картинки, предметы, презентация).",
                        "Беседа-игра (с загадками, физкультминуткой, ролевыми элементами).",
                        "Круглый стол / обсуждение (для старших классов).",
                    ],
                    index=0,
                    key="talk_format",
                    on_change=_talk_params_changed,
                )

                talk_elements = st.multiselect(
                    "Конкретные элементы, которые нужно включить",
                    [
                        "Ключевые термины или имена",
                        "Один конкретный факт или история",
                        "Связь с местным контекстом",
                        "Стихотворение, короткая песня, пословица",
                        "Вопросы к аудитории",
                    ],
                    default=[],
                    key="talk_elements",
                    on_change=_talk_params_changed,
                    placeholder="Выберите параметры",
                )

                talk_time_min = st.slider(
                    "Время беседы (мин)",
                    min_value=5,
                    max_value=40,
                    value=int(st.session_state.get("talk_time_min", 15) or 15),
                    step=5,
                    key="talk_time_min",
                    on_change=_talk_params_changed,
                )
                talk_notes = st.text_area(
                    "Примечания (опционально)",
                    value=st.session_state.get("talk_notes", ""),
                    key="talk_notes",
                    on_change=_talk_params_changed,
                    height=80,
                )

            with right:
                talk_integration = st.multiselect(
                    "Интеграция с другими предметами (межпредметные связи)",
                    [
                        "Чтение (сказка, миф, отрывок из произведения).",
                        "Музыка (песня, звуки природы).",
                        "ИЗО (рассматривание репродукции, описание картины).",
                        "Окружающий мир / история (исторический контекст, природные явления).",
                        "Технология (поделка, которую можно сделать после).",
                    ],
                    default=[],
                    key="talk_integration",
                    on_change=_talk_params_changed,
                    placeholder="Выберите параметры",
                )

                # Отображаем пользователю только короткие метки, а подсказки для промпта
                # храним отдельно и добавляем в auto_prompt (скрыты от учителя)
                audience_hints = {
                    "Класс с высоким уровнем развития": "нужны проблемные вопросы, элементы дискуссии, можно использовать более сложные термины",
                    "Обычный общеобразовательный класс": "нужна опора на наглядность и простые формулировки",
                    "Класс с низкой учебной мотивацией": "тяжело воспринимают информацию на слух. Нужно: очень короткие блоки по 2-3 минуты, постоянная смена деятельности, больше игровых моментов, простой язык, минимум терминов",
                }

                talk_audience = st.selectbox(
                    "Особенности аудитории",
                    [
                        "Класс с высоким уровнем развития",
                        "Обычный общеобразовательный класс",
                        "Класс с низкой учебной мотивацией",
                    ],
                    index=1,
                    key="talk_audience",
                    on_change=_talk_params_changed,
                )

                talk_tone = st.selectbox(
                    "Тон общения",
                    [
                        "—",
                        "Доброжелательный",
                        "Сказочный",
                        "Научно-популярный",
                        "Доверительный",
                        "Энергичный",
                    ],
                    index=0,
                    key="talk_tone",
                    on_change=_talk_params_changed,
                )

            # Авто-промпт для беседы
            def _as_opt(v: str) -> str:
                v = (v or "").strip()
                return "" if v == "—" else v

            goal_txt = _as_opt(st.session_state.get("talk_goal", ""))
            format_txt = _as_opt(st.session_state.get("talk_format", ""))
            audience_txt = _as_opt(st.session_state.get("talk_audience", ""))
            tone_txt = _as_opt(st.session_state.get("talk_tone", ""))

            elements_txt = ", ".join(st.session_state.get("talk_elements") or [])
            integration_txt = ", ".join(st.session_state.get("talk_integration") or [])
            minutes = int(st.session_state.get("talk_time_min") or 15)

            auto_prompt = (
                "Ты — опытный учитель и педагог. Подготовь сценарий беседы для обсуждения с детьми по теме урока.\n\n"
                f"Предмет: {talk_subject}.\n"
                f"Класс: {talk_grade}.\n"
                "ТЕМА (главная):\n"
                f"\"\"\"\n{talk_topic or '(тема не указана)'}\n\"\"\"\n"
                f"Длительность беседы: ~{minutes} минут.\n\n"
            )
            if goal_txt:
                auto_prompt += f"Цель и педагогический фокус: {goal_txt}\n"
            if format_txt:
                auto_prompt += f"Формат и методы: {format_txt}\n"
            if elements_txt:
                auto_prompt += f"Обязательно включить элементы: {elements_txt}.\n"
            if integration_txt:
                auto_prompt += f"Межпредметные связи: {integration_txt}.\n"
            talk_notes_txt = (st.session_state.get("talk_notes") or "").strip()
            if talk_notes_txt:
                auto_prompt += f"Примечания учителя: {talk_notes_txt}\n"
            if audience_txt:
                auto_prompt += f"Особенности аудитории: {audience_txt}\n"
                # Добавляем скрытую подсказку в автоматически формируемый промпт,
                # чтобы учитель её не видел в интерфейсе, но модель учитывала.
                hint = ""
                try:
                    hint = audience_hints.get(st.session_state.get("talk_audience", ""), "")
                except Exception:
                    hint = ""
                if hint:
                    auto_prompt += f"В промпте добавь - {hint}\n"
            if tone_txt:
                auto_prompt += f"Тон общения: {tone_txt}.\n"

            auto_prompt += (
                "\nТребования к результату:\n"
                "- Соответствуй возрасту/классу; без спорных/взрослых тем.\n"
                "- Структура: 1) Вступление (1–2 мин) 2) Рассказ/сюжет 3) Вопросы и наводящие реплики 4) Мини-активность 5) Итог.\n"
                "- Дай текст учителя и вопросы к классу (чтобы дети больше говорили).\n"
                "- Язык простой, живой; без воды.\n"
                "\nВывод: строго в Markdown (без HTML). Не пиши рассуждений — только готовый сценарий беседы.\n"
            )

            st.session_state.setdefault("talk_prompt_autofill", True)
            if st.session_state.get("talk_prompt_autofill"):
                st.session_state["talk_prompt_editor"] = auto_prompt

            if False:
                with st.expander("Шаблон для генерации (можно править)", expanded=False):
                    prompt_cols = st.columns([1, 3])
                    if prompt_cols[0].button("Сбросить к авто", key="talk_prompt_reset_btn"):
                        st.session_state["talk_prompt_autofill"] = True
                        st.session_state["talk_prompt_editor"] = auto_prompt

                    st.text_area(
                        "Промпт",
                        height=240,
                        key="talk_prompt_editor",
                        on_change=lambda: st.session_state.__setitem__("talk_prompt_autofill", False),
                    )

            st.session_state.setdefault("talk_generated_md", "")
            gen_cols = st.columns([1, 1, 2])
            if gen_cols[0].button("Сгенерировать беседу (ИИ)", key="talk_generate_btn"):
                api_key_local = st.session_state.get("api_key") or os.getenv("OPENROUTER_API_KEY")
                if not api_key_local:
                    st.error("Укажите OpenRouter API key в админ-панели (sk-or-v1-...).")
                elif not talk_topic:
                    st.warning("Сначала задайте тему (в форме плана урока сверху).")
                else:
                    prompt_text = (st.session_state.get("talk_prompt_editor") or auto_prompt).strip()
                    with st.spinner("Генерирую беседу..."):
                        resp = generate_with_deepseek(api_key_local, prompt_text)
                    text = (
                        resp.get("choices", [{}])[0]
                        .get("message", {})
                        .get("content", "")
                    )
                    st.session_state["talk_generated_md"] = (text or "").strip()

            if gen_cols[1].button("Очистить", key="talk_clear_btn"):
                st.session_state["talk_generated_md"] = ""

            if (st.session_state.get("talk_generated_md") or "").strip():
                st.markdown("### Результат (можно править)")
                st.session_state["talk_generated_md"] = st.text_area(
                    "talk_generated_md_editor",
                    value=st.session_state["talk_generated_md"],
                    height=360,
                    label_visibility="collapsed",
                )

                # Сохранение результата в .docx
                try:
                    talk_docx_title = _normalize_material_filename(
                        grade=talk_grade,
                        kind="беседа",
                        topic=talk_topic or "беседа",
                        ext="docx",
                    )

                    src_md = st.session_state.get("talk_generated_md") or ""
                    md_norm = normalize_ai_markdown(_postprocess_plan_text(src_md))
                    html_for_docx = quill_html_utils.sanitize_html_for_quill(
                        markdown_to_html(md_norm)
                    )
                    bytes_docx = _html_to_docx_bytes(html_for_docx)

                    btn_cols = st.columns([1, 1, 2])
                    try:
                        btn_cols[0].download_button(
                            "Сохранить docx",
                            data=bytes_docx,
                            file_name=talk_docx_title,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            on_click="ignore",
                            key="talk_download_docx_btn",
                        )
                    except Exception:
                        logging.exception("Error while creating download button for talk")
                        st.error("Не удалось подготовить файл для скачивания. Попробуйте ещё раз.")

                    # Кнопка для сохранения метаданных материала в БД (без пути к файлу)
                    def _register_talk():
                        try:
                            create_material(
                                filename=talk_docx_title,
                                uploader_id=st.session_state.get("user_id"),
                                topics=(talk_topic or "").strip() or None,
                                path=None,
                                subject=talk_subject,
                                grade=talk_grade,
                            )
                            st.success("Беседа сохранена в базе как материал.")
                            st.session_state["talk_last_registered"] = True
                        except Exception:
                            logging.exception("Error while registering talk in DB")
                            st.error("Не удалось сохранить метаданные беседы в БД.")

                    # Скрытая кнопка "Сохранить в БД"
                    if False:
                        _register_talk()
                except Exception as e:
                    logging.exception("Error while converting talk to docx")
                    st.error(f"Не удалось собрать .docx: {e}")

    # Подсказка: быстрый поиск и просмотр материалов (как дополнительный блок)
    # Блок временно скрыт — обернут в условие, чтобы не рендериться в UI.
    if False:
        with st.expander("Поиск по материалам / AI-помощник"):
            query = st.text_input("Введите запрос для поиска по материалам и AI", key="search_query")
            if st.button("Поиск", key="search_btn"):
                if not query:
                    st.warning("Введите запрос.")
                else:
                    with st.spinner("🔍 Ищу..."):
                        results = []
                        if api_key:
                            try:
                                headers = {
                                    "Authorization": f"Bearer {api_key}",
                                    "Content-Type": "application/json",
                                    "HTTP-Referer": "https://teacher-assistant.streamlit.app",
                                    "X-Title": "Teacher Assistant Search"
                                }
                                data = {
                                    "model": "deepseek/deepseek-chat",
                                    "messages": [{"role": "user", "content": f"Помоги учителю найти: {query}"}],
                                    "temperature": 0.5,
                                    "max_tokens": 800,
                                }
                                resp = requests.post(DEEPSEEK_API_BASE, headers=headers, json=data, timeout=30)
                                resp.raise_for_status()
                                ai_response = resp.json()["choices"][0]["message"]["content"]
                                results.append({"title": f"AI: {query}", "snippet": ai_response, "type": "ai_response"})
                            except Exception:
                                pass

                        for p in materials_dir.glob("**/*"):
                            if query.lower() in p.name.lower():
                                results.append({"title": f"Файл: {p.name}", "snippet": str(p), "type": "local_file"})

                    if results:
                        for r in results:
                            with st.expander(r["title"]):
                                if r["type"] == "ai_response":
                                    st.markdown(r["snippet"])
                                else:
                                    st.write(r["snippet"])
                    else:
                        st.info("Ничего не найдено.")

elif app_mode == "Админ-панель":
    st.header("Админ-панель")
    if not is_admin():
        st.info("Введите пароль администратора в сайдбаре, чтобы открыть админ-панель.")
    else:
        st.subheader("Загрузка материалов")
        st.caption("Сначала выберите предмет и класс, затем загрузите файл КТП/материала.")
        upload_subject = st.selectbox("Предмет для загружаемых материалов", SUBJECTS, index=0, key="upload_subject")
        upload_grade = st.selectbox("Класс для загружаемых материалов (класс)", GRADES, index=3, key="upload_grade")
        uploaded_files = st.file_uploader(
            "Загрузите конспекты или раздаточные материалы",
            accept_multiple_files=True,
        )
        if uploaded_files:
            for f in uploaded_files:
                save_path = materials_dir / f.name
                with open(save_path, "wb") as out:
                    out.write(f.getbuffer())


                # Попытка извлечь темы из файла с помощью общего парсера
                suffix = save_path.suffix.lower()
                suggestions = []
                # Для PDF даём администратору возможность указать номер столбца
                if suffix == ".pdf":
                    col_key = f"{f.name}_pdf_col"
                    col = st.number_input("Номер колонки для извлечения тем (1-based)", min_value=1, max_value=10, value=2, key=col_key)
                    include_text_key = f"{f.name}_include_text"
                    include_text = st.checkbox("Включать извлечение из свободного текста (много шума)", value=False, key=include_text_key)
                    if st.button("Извлечь темы из PDF (по колонке)", key=f"extract_pdf_{f.name}"):
                        try:
                            suggestions = extract_topics_from_pdf_advanced(save_path, max_topics=500, column=col, include_text=include_text)
                            st.success(f"Найдено {len(suggestions)} темы(а) в PDF.")
                        except Exception as e:
                            st.error(f"Ошибка извлечения из PDF: {e}")
                # Если не PDF или не нажата кнопка — используем общий парсер
                if not suggestions:
                    suggestions = extract_topics(save_path, max_topics=50)

                # Извлекаем весь текст файла построчно для ручной правки
                full_lines = extract_full_text(save_path)

                st.write(f"Файл сохранён: {save_path}")
                st.caption("Отредактируйте строки и отметьте те, которые являются темами")

                # Список тем из существующих планов для выбора (выпадашки скрываем, если пусто)
                existing_plans = list_lesson_plans(limit=200)
                existing_topics = sorted({p.topic for p in existing_plans if p.topic})

                # Инициализация session_state для строк и чекбоксов
                for i, line in enumerate(full_lines):
                    key_text = f"{f.name}_text_{i}"
                    key_pick = f"{f.name}_pick_{i}"
                    if key_text not in st.session_state:
                        st.session_state[key_text] = line
                    # помечаем как предложенную тему, если строка похожа на suggestion
                    initial_pick = any(s.strip().lower() in line.strip().lower() for s in suggestions)
                    if key_pick not in st.session_state:
                        st.session_state[key_pick] = initial_pick

                # Инициализация session_state для тем, извлечённых из PDF (если есть)
                if suggestions:
                    st.session_state.setdefault(f"{f.name}_pdf_count", len(suggestions))
                    for i, t in enumerate(suggestions):
                        key_text = f"{f.name}_pdf_text_{i}"
                        key_pick = f"{f.name}_pdf_pick_{i}"
                        if key_text not in st.session_state:
                            st.session_state[key_text] = t
                        if key_pick not in st.session_state:
                            st.session_state[key_pick] = True

                with st.expander("Просмотр и разметка строк (нажмите, чтобы открыть)"):
                    for i, line in enumerate(full_lines):
                        key_text = f"{f.name}_text_{i}"
                        key_pick = f"{f.name}_pick_{i}"
                        col1, col2 = st.columns([0.08, 0.92])
                        with col1:
                            st.checkbox("", value=st.session_state.get(key_pick, False), key=key_pick)
                        with col2:
                            st.text_input(f"Строка {i+1}", key=key_text)

                # Отдельный блок для тем, извлечённых из PDF (если есть)
                if suggestions:
                    with st.expander("Темы, извлечённые из файла (проверить/править)"):
                        for i in range(st.session_state.get(f"{f.name}_pdf_count", 0)):
                            key_text = f"{f.name}_pdf_text_{i}"
                            key_pick = f"{f.name}_pdf_pick_{i}"
                            col1, col2 = st.columns([0.08, 0.92])
                            with col1:
                                st.checkbox("", value=st.session_state.get(key_pick, False), key=key_pick)
                            with col2:
                                st.text_input(f"Тема (из PDF) {i+1}", key=key_text)

                # Дополнительные возможности: добавить пустую строку как тему
                if st.button(f"Добавить пустую строку для {f.name}"):
                    idx = len(full_lines)
                    key_text = f"{f.name}_text_{idx}"
                    key_pick = f"{f.name}_pick_{idx}"
                    st.session_state[key_text] = ""
                    st.session_state[key_pick] = True
                    safe_rerun()

                if st.button(f"Сохранить метаданные для {f.name}"):
                    from storage import create_material

                    picks = []
                    for i in range(len(full_lines)):
                        key_text = f"{f.name}_text_{i}"
                        key_pick = f"{f.name}_pick_{i}"
                        if st.session_state.get(key_pick):
                            txt = st.session_state.get(key_text, "").strip()
                            if txt:
                                picks.append(txt)

                    topics_csv = ','.join(dict.fromkeys(picks)) if picks else None
                    user_id = st.session_state.get('user_id')
                    create_material(
                        filename=f.name,
                        uploader_id=user_id,
                        topics=topics_csv,
                        path=str(save_path),
                        subject=upload_subject,
                        grade=upload_grade,
                    )
                    st.success("Метаданные материала сохранены.")

            st.success(f"Сохранено {len(uploaded_files)} файл(ов) в папку {materials_dir}/")

        st.subheader("Материалы (обзор и управление)")
        topic_filter = st.text_input("Фильтр по темам материалов (подстрока)")
        materials = list_materials(limit=200)
        if not materials:
            st.info("Пока нет сохранённых метаданных материалов. Загрузите файлы выше или используйте bulk_upload.")
        else:
            for m in materials:
                if topic_filter and (not m.topics or topic_filter.lower() not in m.topics.lower()):
                    continue
                st.markdown(f"**{m.filename}**")
                st.caption(f"Темы: {m.topics or '—'}")
                if m.path and Path(m.path).exists():
                    with open(m.path, "rb") as fh:
                        try:
                            st.download_button(
                                label="Скачать файл",
                                data=fh.read(),
                                file_name=Path(m.path).name,
                                key=f"material_download_{m.id}",
                                on_click="ignore",
                            )
                        except KeyError:
                            logging.exception("Missing media key when creating material download button")
                            st.error("Файл временно недоступен. Пожалуйста, перезагрузите страницу и повторите попытку.")
                        except Exception:
                            logging.exception("Error while creating material download button")
                            st.error("Не удалось подготовить файл для скачивания. Попробуйте ещё раз.")

                if st.button("Удалить материал", key=f"delete_material_{m.id}"):
                    if m.path and Path(m.path).exists():
                        try:
                            os.remove(m.path)
                        except Exception:
                            pass

                    delete_material(m.id)
                    st.success("Материал удалён.")
                    safe_rerun()

                st.markdown("---")

st.caption(
    "Разработано учителем информатики Ларченко А.П. ГУО 'Средняя школа №16 г. Минска'.\n"
    "\nРабота ведётся в рамках проекта по цифровизации образовательного процесса. \n2026 год."
)
